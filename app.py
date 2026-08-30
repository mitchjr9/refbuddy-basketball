"""
RefBuddy — Your Minnesota HS Basketball Referee Assistant & Film Coach
Version 1.1 — Feature parity with RefBuddy Football v1.2

Changes from v1.0:
  - PASSWORD GATE: get_secret() reads OS env → secrets.toml → legacy nested table,
    so the app works on Render (env vars) and locally (secrets.toml) alike. Every
    st.secrets access is wrapped in a broad `except Exception` because touching
    st.secrets with no secrets.toml on disk raises StreamlitSecretNotFoundError,
    which is NOT a KeyError and will crash the app on boot if caught too narrowly.
    Fail-closed: missing APP_PASSWORD on Render refuses to start.
  - SESSION USAGE CAP: MAX_FRAMES_PER_SESSION = 400 caps how many analysis frames
    one browser session can send. Sidebar shows the remaining budget.
  - PROMPT CACHING: CORE_KNOWLEDGE is sent as a cached system block on every call
    (identical prefix across the whole app), cutting repeat input cost ~90%.
  - FILM & GRADE: the old Game Film and RefGrade tabs are merged into one tab that
    accepts PHOTOS (jpg/jpeg/png) as well as video (mp4/mov). Stills convert via
    image_to_frame_b64() into the same base64-JPEG format extract_frames() emits,
    so photos and video frames share one identical downstream path. A radio picks
    "Ask a Question" (SYSTEM_PROMPT) or "RefGrade Evaluation" (REFGRADE_PROMPT).
  - REF HUB: Crew Eval and Ref Eval were near-identical forms differing only in
    scope, so they merge into one Evaluations section with a scope dropdown.
    Pre-Game Meeting leads because it is most used and needs no upload.
  - BRANDING: "Powered by Claude" lockup via _asset_data_uri("Claude.png"), with a
    text fallback if the file is missing. "Built by a ref, for refs" everywhere.
  - LEGAL: Terms of Use, non-affiliation notice, and no-warranty disclaimer.

Tabs: 🏀 Home | 🎬 Film & Grade | 👥 Ref Hub | 📝 Quiz
Run:  streamlit run app.py

Repo root must also contain: Claude.png, requirements.txt, .streamlit/config.toml
"""

# ── Standard library ──────────────────────────────────────────────────────────
import base64
import datetime
import functools
import hmac
import json
import os
import subprocess
import sys
import tempfile
import urllib.parse

# ── Third-party ───────────────────────────────────────────────────────────────
import anthropic
import numpy as np
import streamlit as st

# ── OpenCV — auto-install if missing ─────────────────────────────────────────
try:
    import cv2
    OPENCV_AVAILABLE = True
except ImportError:
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "opencv-python-headless", "-q"]
        )
        import cv2
        OPENCV_AVAILABLE = True
    except Exception:
        OPENCV_AVAILABLE = False


# =============================================================================
# CORE KNOWLEDGE BASE
# =============================================================================

CORE_KNOWLEDGE = """
# RefBuddy Core Knowledge Base
## Minnesota High School Basketball — Referee Reference

---

## 0. 2023–2026 NFHS & MSHSL RULES CHANGES & UPDATES

> **INSTRUCTION FOR ALL RESPONSES:** Default to the `2023-2024_NFHS_Basketball_Rulebook.md` baseline (Sections 1–10 below) for any question unless a specific update in THIS section overrides it. Always cite the year when a change applies. If a rule changed in 2023-24 and was further updated in 2025-26, note both years.

---

### 2023-24 Rule Changes (Baseline Season)

**[2023-24] Bonus Free Throw System Revamped — Rule 4-8-1**
- **Old:** One-and-one bonus beginning with the 7th team foul in the half; two-shot bonus on the 10th.
- **New:** NO one-and-one. Two free throws (bonus) beginning with the team's **5th foul in each quarter**. Team fouls reset to zero at the end of each quarter.
- **MSHSL EXCEPTION:** MSHSL retains the old half-based system — 1&1 on 7th team foul per half; two-shot bonus on 10th foul per half (Minnesota Modification K).
- **Why it matters:** HUGE divergence. In NFHS-only games (non-MSHSL), the quarter-reset applies. In ALL MSHSL varsity games, use the half-based 1&1/double-bonus system. Always confirm which system applies at pregame.

**[2023-24] Four Designated Throw-In Spots — Rules 4-36, 6-4-3, 7-5-2 through 4**
- **Old:** General throw-in spots with less specific guidance on frontcourt retentions.
- **New:** When a team retains or gains team control in its frontcourt (due to a violation, common foul prior to bonus, or other stoppages other than out-of-bounds), four designated spots apply: the nearest 28-foot mark on each sideline OR the nearest spot 3 feet outside the lane line on the end line.
- **Updated 2025-26:** The 3-point line now serves as the line of demarcation. Inside the arc = end-line spot (3 ft outside lane); outside the arc = 28-foot sideline spot. Uses visible court markings instead of imaginary lines.
- **Why it matters:** Determines where you hand/bounce the ball for most non-OOB stoppages in the frontcourt. Know your court markings.

**[2023-24] Like-Colored Uniform Bottoms — Rule 3-4-5**
- **New:** Teammates must all wear like-colored uniform bottoms but may wear multiple styles.
- **Why it matters:** Equipment check — all bottoms same color even if different cut/style.

**[2023-24] Undershirt Color — Visiting Team — Rule 3-5-6**
- **New:** Visiting team may wear undershirts that are black OR a single solid color similar to the jersey torso. All teammates must match.
- **Why it matters:** Common uniform question. Visitor can choose black or team color but all must wear the same.

**[2023-24] Shot Clock Operator Location — Rule 2-1-3 NOTE**
- **New:** Shot clock operator shall be seated at the scorer's and timer's table.
- **Why it matters:** Pregame — confirm operator is at the table, not elsewhere.

**[2023-24] Throw-In Correction Window — Rule 7-6-6**
- **New:** An official who administers a throw-in to the wrong team may correct the mistake before the first dead ball after the ball becomes live, unless there is a change of possession.
- **Why it matters:** You can fix a wrong throw-in team without a foul or technical — but only before first dead ball or change of possession.

**[2023-24] Stepping Out of Bounds — Rule 9-3-3**
- **New:** A player may step out of bounds without penalty unless: (a) they are the first to touch the ball after returning to court, or (b) they left the court to avoid a violation.
- **Why it matters:** Clarifies that incidental OOB stepping (e.g., player pushed OOB) is not automatically a violation. The violation triggers only if that player touches the ball first.

---

### 2024-25 Key Points & Mechanics Updates

**[2024-25] Flopping — New Warning System (MSHSL Adopted)**
- **Old:** Faking being fouled was a player technical foul.
- **New:** First team flopping = team WARNING (no technical). Second and subsequent offenses by SAME TEAM = team technical foul (2 FTs + ball at division line).
- **Mechanics:** Show flop signal (#23) during live play, note the time. If offense is advancing, let play continue; issue warning at next dead ball. Do NOT stop play if offense is moving toward basket on a defensive flop.
- **Key point:** Flopping warning doesn't require player to fall — bobbing head, flailing arms counts.
- **Why it matters:** Don't confuse with the five delay-of-game warnings. Flopping warning is a separate category.

**[2024-25] Throw-In Spot — 3-Point Line as Demarcation (Mechanics)**
- Officials use the three-point line to determine frontcourt throw-in spots.
- Inside 3-pt line → end line spot (3 ft outside lane).
- Outside 3-pt line → 28-foot sideline mark.
- Back-court stoppages → four designated spots based on ball location relative to 3-pt line (backcourt tick marks added).

**[2024-25] Uniform Number Contrast (MSHSL memo 12/20/24)**
- Numbers on dark uniforms must clearly contrast. Best practice: white numbers on dark jersey. MSHSL memo from Lisa Quednow, Phil Archer, Jason Nickleby (12/20/24).
- Do NOT stop game for uniform issues — address at next dead-clock opportunity. Repeat issues = contest report.

---

### 2025-26 Rule Changes (Current Season)

**[2025-26] No Offensive Goaltending — Rules 4-22-1 & 4-22-2 (MAJOR)**
- **Old:** Offensive goaltending was a violation.
- **New:** Offensive goaltending violations are ELIMINATED. Only defensive players can commit goaltending.
- **Why it matters:** HUGE change. If an offensive player tips/touches the ball on a downward flight that could enter the basket — NO call. No need to judge "try vs. pass" for offensive team.

**[2025-26] Ball Contacting Backboard = Downward Flight — Rule 4-22-3 (NEW)**
- **New:** Once the ball contacts the backboard, it is automatically considered to be on its downward flight. If a player then touches the ball that has contacted the backboard and has a possibility of entering the basket = GOALTENDING.
- **Why it matters:** Clarifies a common misconception. After backboard contact, always treat the ball as being in downward flight for goaltending purposes.

**[2025-26] Basket Interference — Backboard Slap — Rules 4-6-1a & 4-6-1b (NEW)**
- **New:** Basket interference NOW INCLUDES: a player slapping or striking the backboard, causing it to vibrate, while the ball is on/within the basket, touching the backboard, or within the cylinder.
- **Why it matters:** Formerly just a technical foul. Now it's ALSO basket interference — the points count if the ball would have gone in.
- **Situation:** B1 slaps backboard → basket vibrates while ball is in basket → TWO violations: (1) basket interference = 2 pts awarded, (2) technical foul = 2 FTs + ball at division line. Both are penalized.

**[2025-26] Player Definition Updated — Rule 4-34-1**
- **New:** A player is one of the five team members legally on the court at any time, EXCEPT during time-outs or intermissions.
- **Why it matters:** During a timeout, all players are bench personnel. Technical fouls on bench personnel during timeouts = indirect tech to head coach (loss of coaching box privileges).

**[2025-26] Throw-In Spot — 3-Point Line Demarcation — Rule 7-5-4**
- **New:** The three-point line (visible court marking) now determines the designated throw-in spot following a stoppage in the frontcourt (not due to OOB). Inside = end line. Outside = 28-foot line.
- **Why it matters:** Uses visible markings; eliminates the imaginary line confusion from the prior system.

**[2025-26] Thrower Delay OOB — Now a Violation — Rules 9-2-12 & 9-3-4 (NEW)**
- **Old:** Thrower who stepped out of bounds to deceive and then touched ball = technical foul.
- **New:** This is now a VIOLATION (turnover), not a technical foul.
- **Why it matters:** Lesser penalty, easier to enforce consistently. Player cannot step out of bounds as a thrower to deceive and be first to touch the ball back in bounds.

**[2025-26] Backboard Contact Technical — Rule 10-4-4b**
- **New:** Players may not illegally contact the backboard or ring in ways that create unfair advantage or interfere with a scoring attempt. Technical foul.
- **Why it matters:** Closes gap between basket interference and technical fouls; removes subjective intent standard.

**2025-26 Points of Emphasis:**
1. **Contact on the Ball Handler** — Hand-checking, body displacement (hips/torso), and impeding freedom of movement are fouls. NOT all late-game fouls are intentional. Determine by the ACT, not the coach's verbal instructions ("foul," "red," "scramble"). Consistent standards throughout entire game.
2. **Bench Decorum, Communication, and Player Altercations** — Apply Rule 4-48 (Warning for Coach/Team Conduct) early. Acknowledge reasonable coach inquiries (nod is sufficient). Rules-based questions deserve a clear informative response.
3. **Faking Being Fouled (Flopping)** — First offense: team warning. Subsequent offense: team technical. Don't stop advancing play on a defensive flop.

---

### Quick-Reference: Rules That CHANGED from 2023-24 Baseline

| Rule | Year | What Changed |
|------|------|-------------|
| 4-8-1 | 2023-24 | Bonus = 5th team foul/quarter (2 FTs); no 1&1 [NFHS only — MSHSL retains 1&1 half system] |
| 4-36, 7-5-2–4 | 2023-24 | Four designated frontcourt throw-in spots established |
| 3-4-5 | 2023-24 | Like-colored uniform bottoms required |
| 3-5-6 | 2023-24 | Visiting undershirt: black or jersey color (team must match) |
| 7-6-6 | 2023-24 | Wrong throw-in team correctable before first dead ball / change of possession |
| 9-3-3 | 2023-24 | Stepping OOB not a violation unless first to touch ball on return |
| 2-1-3 NOTE | 2023-24 | Shot clock operator must sit at scorer's/timer's table |
| Flopping rule | 2024-25 | Player tech → team warning (1st); team tech (2nd+) |
| 7-5-4 | 2024-25 / 2025-26 | 3-pt line as demarcation for frontcourt throw-in spot |
| 4-22-1 & 2 | 2025-26 | Offensive goaltending ELIMINATED |
| 4-22-3 | 2025-26 | Ball touching backboard = automatically on downward flight |
| 4-6-1a/b | 2025-26 | Backboard slap causing vibration = basket interference |
| 4-34-1 | 2025-26 | Players become bench personnel during time-outs |
| 9-2-12, 9-3-4 | 2025-26 | Thrower delay OOB = violation (was technical foul) |
| 10-4-4b | 2025-26 | Backboard/ring contact creating unfair advantage = technical foul |

---

---

## 1. NFHS RULE HIERARCHY & KEY CITATIONS

### Rule 1 — The Court and Equipment
- **Court:** 84' × 50' recommended (max 94' × 50'). Division line bisects court (Rule 1-3).
- **Basket height:** 10 feet above floor; ring 18" inside diameter; net 15–18" long (Rule 1-10, 1-11).
- **Ball — Boys:** 29.5–30" circumference, 20–22 oz (Rule 1-12).
- **Ball — Girls:** 28.5–29" circumference, 18–20 oz (Rule 1-12).
- **NFHS Authenticating Mark** required on game balls (Rule 1-12-1g).
- **Backboard padding required** on bottom/sides up to 15" from bottom (Rule 1-9).
- **Coaching box:** 28-foot maximum, state option (Rule 1-13-2). **MSHSL: 14 feet** (MN Mod I).
- **Officials arrive:** Minimum 15 minutes before scheduled start (Rule 2-2-2).
- **Officials' uniform:** Black-and-white striped shirt, black pants, predominantly black shoes/socks (Rule 2-1-1). **MSHSL: Smitty gray shirt with black panel, MSHSL logo on left crest** (MN Mod J).

### Rule 2 — Officials and Their Duties
- **Referee's pregame:** Inspect/approve equipment, designate timepiece, official scorebook, notify teams 3 min before each half, verify coaches confirm legal uniforms (Rule 2-4).
- **No replay equipment** — officials may not use video/replay to make decisions during game (Rule 2-2-1). EXCEPTION: State championship series may permit replay for last-second shot at 0:00.
- **Referee has final authority** on any matter not specifically covered in rules (Rule 2-3).
- **3-person crew:** Referee (R), Umpire 1 (U1/Trail), Umpire 2 (U2/Center or Lead).
  - Lead (L) = under basket official; Trail (T) = half-court; Center (C) = middle position.

### Rule 3 — Uniforms and Equipment
- **Jersey numbers:** Must be 0–5, 10–15, 20–25, 30–35, 40–45, 50–55 (numerals that can be indicated with one hand). No duplicates within a team (Rule 3-4-3).
- **Jersey colors:** Home = white; Visitor = dark (Rule 3-4-1). **MSHSL Mod B:** Home = dark; Visitor = white (opposite of NFHS!).
- **Number contrast (MSHSL memo 12/20/24):** Numbers on dark uniforms must clearly contrast. Best practice = white numbers.
- **Undershirts (home):** Must be white, hemmed (Rule 3-5-1). **Undershirts (visitor):** Single solid color similar to jersey torso OR black; all teammates must match (Rule 3-5-6 [2023-24 change]).
- **Compression apparel:** Black, white, beige, or predominant jersey color; all teammates same color. **MSHSL Mod C:** All apparel (wristbands, headbands, sleeves, knee pads, compression) must be solid black OR white; same color for all participants.
- **Jewelry prohibited** (Rule 3-5-3). Medical alert jewelry may be taped.
- **Headbands/wristbands:** Single solid color, must match teammates (Rule 3-5-3c). MSHSL: must be black or white.

### Rule 4 — Definitions
- **Backcourt/Frontcourt:** Team's frontcourt = between its end line and nearer edge of division line, including its basket (Rule 4-13).
- **Bonus:** 2 FTs awarded for common foul beginning with 5th team foul per quarter (NFHS Rule 4-8-1). **MSHSL: 1&1 on 7th foul per half; double bonus on 10th** (MN Mod K).
- **Closely Guarded:** Opponent within 6 feet in player's frontcourt, holding or dribbling the ball. 5-second count applies (Rule 4-10). **MSHSL EXCEPTION (Mod F):** Closely guarded is NOT in effect when player is DRIBBLING the ball (only applies to player holding the ball).
- **Control (Player):** Holding or dribbling a live ball (Rule 4-12-1).
- **Control (Team):** While player controls, live ball passed among teammates, interrupted dribble, or player has disposal for throw-in (Rule 4-12-2).
- **Dribble:** Begins when player pushes/throws/bats ball to floor before pivot foot is lifted; ends when ball rests in hand(s), both hands touch simultaneously, or ball becomes dead (Rule 4-15).
- **Foul types:** Common foul (Rule 4-19-2), intentional foul (Rule 4-19-3), flagrant foul (Rule 4-19-4), technical foul (Rule 4-19-5), player-control foul (Rule 4-19-6), team-control foul (Rule 4-19-7).
- **Goaltending [2025-26 UPDATED]:** (a) Defensive player touches ball during try/tap while ball is in downward flight entirely above ring level with possibility of entering basket, not touching cylinder (Rule 4-22-1). (b) Ball contacting backboard = automatically on downward flight (Rule 4-22-3 NEW). NO more offensive goaltending (Rule 4-22-1 & 2 [2025-26]).
- **Guarding (legal):** Both feet on floor, torso facing opponent. Max 6 feet in closely guarded (Rule 4-23).
- **Held ball:** Opponents' hands firmly on ball so control can't be obtained without undue roughness (Rule 4-25).
- **Pivot:** Stepping with same foot while other (pivot) foot kept at point of contact (Rule 4-33).
- **Player [2025-26 UPDATED]:** One of five team members legally on court; becomes bench personnel during time-outs and intermissions (Rule 4-34-1).
- **Secondary Defender (MSHSL restricted area):** Teammate who helped a primary defender beaten by offensive player (head/shoulders past defender); or double-teams low post; or is outnumbered in fast break (MSHSL addendum Rule 4-41).
- **Restricted Area (MSHSL addendum Rule 4-38):** 4-foot radius arc from center of basket to inside of arc line, extending to face of backboard. Secondary defender "in" area when any part of either foot is in or above this area.
- **Traveling:** Moving pivot foot or taking more than allowed steps while holding ball (Rule 4-44 / 9-4-1).

### Rule 5 — Scoring and Timing
- **Goal:** Ball enters basket from above and remains/passes through net; counts for team into whose basket it falls regardless of who threw it (Rule 5-1).
- **Field goal value:** 2 points; 3 points if from beyond 3-point line (Rule 5-1-2,3).
- **Free throw:** 1 point (Rule 5-1-4).
- **Game length:** Four 8-minute quarters (Rule 5-4). **MSHSL Mod A:** Two 18-minute halves (varsity); max 16-minute halves below varsity.
- **Extra period:** 4 minutes for overtime (Rule 5-7).
- **Game clock stops:** On every foul, violation, held ball, OOB, time-out, etc. (Rule 5-8).
- **Time-outs (Rule 5-11-1 — CONFIRMED, full text in Section 11):** three 60-second and two 30-second time-outs per team per regulation game; one additional 60-second time-out per extra period; unused time-outs accumulate and may be used at any time. Excess time-outs are granted but cost a technical foul (5-11-6, 10-2-3).
- **Mercy Rule (MSHSL Mod G):** When point differential reaches **35 or more** with **less than 9 minutes remaining** in second half → running clock. Clock stops only for time-outs. Shot clock continues. Returns to regular timing if differential drops to **less than 30**.

### Rule 6 — Live Ball/Dead Ball
- **Ball becomes live:** On a throw-in, when the ball is released; on a jump ball, when the ball leaves the official's hand; on a free throw, when at the disposal of the free thrower (Rule 6-1).
- **Ball becomes dead:** On any foul, violation, held ball, time-out, successful goal, when period expires (Rule 6-7).
- **Resumption-of-play procedure:** After certain violations — no substitute, no time-out, no delay (Rule 4-38).

### Rule 7 — Out of Bounds and the Throw-In
- **Out of bounds:** When ball or player touches the boundary line, floor, or objects on/outside it (Rule 7-1).
- **Responsibility:** Last player to touch ball before it goes OOB is responsible (Rule 7-2).
- **Throw-in administration (TRAIL):** Bounce pass or direct pass; 5-second count; thrower may not leave spot until ball is released (Rule 7-6).
- **Four designated throw-in spots (frontcourt, 2025-26 updated Rule 7-5-4):**
  - **Inside 3-pt line** (stoppage occurred inside 3-pt line or on 3-pt line) → end line, 3 feet outside lane line (nearest to stoppage).
  - **Outside 3-pt line** (stoppage occurred outside arc) → 28-foot sideline mark nearest to stoppage.
  - These apply when team retains/gains possession in frontcourt due to violation, common foul (pre-bonus), or stoppages other than OOB.
- **Throw-in violations (Rule 7-6-7):** Leaves spot, 5-second count, passes through basket without touching player, steps on/over boundary before released, etc.
- **Thrower delay OOB [2025-26 NEW Rule 9-2-12 & 9-3-4]:** Purposely stepping OOB as thrower then being first to touch ball in bounds = VIOLATION (was technical foul).
- **Wrong team throw-in correction (Rule 7-6-6):** Can correct before first dead ball / change of possession.

### Rule 8 — Free Throw
- **Free thrower:** In semicircle, behind free-throw line. 10-second limit after ball at disposal (Rule 8-3).
- **Free thrower violations (Rule 9-1):** Cross line before ball touches ring/board/basket = ball dead on violation; opponent's violation = ball live (attempt counts if successful) or retake.
- **Lane occupancy:** Players take assigned lane spaces; alternating positions within 3 feet of lane; may not enter until ball is released by shooter (Rule 8-1, 8-2).
- **Penalty:** Free throw violation by shooter = ball dead, no point; free throw violation by non-shooter on defense = if made it counts, if missed try again (Rule 9-1 penalties).
- **Technical foul free throws:** 2 FTs, no lane players, ball inbounded at division line opposite scorer's table after last FT.

### Rule 9 — Violations and Penalties
- **Throw-in violations (Rule 9-2):** 5-second count, leave designated spot, enter court before released, thrower-in contacted (PENALTY: intentional foul on contact with thrower-in), ball passed through basket directly.
- **Out of bounds (Rule 9-3):** Ball awarded to opponents. Player may step OOB without penalty unless first to touch ball on return or left court to avoid violation (Rule 9-3-3).
- **Traveling (Rule 9-4):** Pivot foot lifted and replaced, two steps without dribble, jumping with ball and not releasing before touching floor.
- **Illegal dribble (Rule 9-5):** Dribble with two hands, carry/palm, dribble a second time after dribble ends (double dribble).
- **Three seconds (Rule 9-7):** Player in their team's frontcourt may not remain in the lane for more than 3 consecutive seconds while team is in control and clock is running.
- **Ten seconds (backcourt) (Rule 9-8):** Team must advance from backcourt to frontcourt within 10 seconds. **MSHSL:** Use shot clock for 10-second count when shot clock is operating.
- **Backcourt violation (Rule 9-9):** Once ball established in frontcourt, team in control may not return to backcourt. Note: if offense last touches in frontcourt and first touches in backcourt = violation.
- **Closely guarded (Rule 9-10):** Holding ball in own frontcourt while closely guarded for 5 consecutive seconds = violation. **MSHSL Mod F:** NOT in effect while DRIBBLING.
- **Goaltending/Basket interference (Rule 9-11 / 4-6, 4-22 [2025-26]):**
  - Goaltending: defensive player touches ball on downward flight above ring level (or after backboard contact per new Rule 4-22-3) = 2 pts awarded.
  - Basket interference: touching ball/basket while ball is in/on basket, or within cylinder; also backboard slap causing vibration while ball in/on basket (Rule 4-6-1a/b NEW 2025-26) = 2 pts awarded.
  - Offensive goaltending ELIMINATED (2025-26).
- **Kicking/fisting ball (Rule 9-4-3):** Intentional striking = violation (ball to opponents).
- **Excessive swinging of arms/elbows (Rule 9-12):** Violation if it endangers opponent.

### Rule 10 — Fouls and Penalties
- **Administrative technical foul (Rule 10-1):** Roster errors, equipment violations, delay.
- **Team technical foul (Rule 10-2):** Delay of game after warning, too many players, improperly worn equipment after warning.
- **Player technical foul (Rule 10-4):** Unsporting act, delay, illegal equipment, taunting, hanging on rim, backboard slap (Rule 10-4-4b [2025-26]).
- **Bench technical foul (Rule 10-5):** Coach, substitute, team attendant, or follower commits unsporting acts. Bench tech = indirect tech to head coach = loss of coaching box.
- **Head Coach's Rule (Rule 10-6):**
  - Direct tech: assessed directly to head coach for abusive conduct. Coach gets 2 FTs against them; flagrant or second direct = ejection.
  - Indirect tech: bench tech charged to bench member → indirect tech charged to head coach. First indirect = warning, loss of coaching box. Second indirect in same half = ejection of head coach.
- **Personal foul penalty:** Offended player shoots FTs (if in bonus), or throw-in if not in bonus.
- **Intentional foul penalty:** 2 FTs + throw-in at point of interruption regardless of bonus.
- **Flagrant foul:** 2 FTs + throw-in; offender ejected (Rule 10-7-12).
- **Player disqualification:** 5th foul (personal + technical fouls combined); 2 tech fouls; 1 flagrant foul (Rule 4-14).
- **Technical foul penalty:** 2 FTs; ball at division line opposite scorer's table. Counts as team foul.

---

## 2. MSHSL MINNESOTA-SPECIFIC RULES & MODIFICATIONS

### Key Minnesota Modifications (December 2025)
- **Mod A — Game Length (Varsity):** Two 18-minute halves (not quarters). JV/sub-varsity: max 16-minute halves.
- **Mod B — Uniform Colors:** HOME = dark uniforms; VISITOR = white uniforms. (OPPOSITE of NFHS default.)
- **Mod C — Apparel Colors:** All apparel (excluding knee braces) — wristbands, headbands, arm/knee sleeves, knee pads, compression shorts, tights — must be solid **black or white**, same color for all participants.
- **Mod D — Sub-Varsity:** Also played in halves, max 16 minutes.
- **Mod E — Shot Clock:** 35-second shot clock mandatory for ALL varsity games, including extra periods. (Adopted 2023-24.)
- **Mod F — Closely Guarded Exception:** Closely guarded rule (5-second count) is NOT in effect when player is dribbling the ball (Rule 9-10-1a NOTE). Applies only to player holding the ball.
- **Mod G — Mercy Rule:** Point differential **35+** with **<9 minutes remaining** in the second half → running clock (game clock only stops for time-outs). Returns to regular timing if differential drops to **<30 points**. Shot clock continues throughout.
- **Mod H — Restricted Area Arc:** Required for all levels of MSHSL play. 4-foot radius from center of basket.
- **Mod I — Coaching Box:** 14 feet (NFHS maximum is 28 feet).
- **Mod J — Officials Uniform:** Smitty gray shirt with black panel, black collar, sleeve cuffs, MSHSL logo on left crest. All crew members must match.
- **Mod K — Bonus Free Throws:** 1&1 on 7th team foul per HALF; double bonus (2 FTs) on 10th team foul per half. (MSHSL does NOT use the NFHS quarter-reset system.)

---

## 3. MSHSL SHOT CLOCK — COMPLETE RULES (2025-26)

### Basics
- **Shot clock period:** 35 seconds (standard); 20 seconds (offensive rebound situations).
- **Mandatory:** All varsity games including extra periods; mercy rule games.
- **Turn off:** When game clock ≤ shot clock period (i.e., game clock ≤ 35 seconds).
- **Operator:** Seated at scorer's/timer's table. Separate distinct horn. Must have backup stopwatch and air horn.
- **Both clocks must work** to start game. If one fails → turn off both, use alternate procedure.
- **Shot clock horns do NOT stop play.** Only a whistle indicating a violation stops play.
- **Recall function:** Not required but strongly recommended.

### When Shot Clock STARTS
- Any throw-in when ball is legally touched/touches any player on court (does not need to be in possession).
- Jump ball: when ball is POSSESSED (not when it's tipped).
- Change in team control.

### RESET to 35 Seconds (Full Reset)
1. Scored basket (starts when ball legally touched after throw-in).
2. Change in team control while ball remains live.
3. Single personal foul in the BACKCOURT.
4. Kicked or fisted ball by defense in the BACKCOURT.
5. Free throw situation: set to :35 immediately; offense rebounds → reset to :20.
6. Defense causes held ball during team control; AP arrow favors defense.

### RESET to 20 Seconds (Offensive Rebound Situations)
1. Offense gains control anywhere after unsuccessful field goal that contacts ring/flange.
2. Offense gains control anywhere after unsuccessful free throw remaining in play.
3. Defense fouls after FG miss (hits rim) or during successful try, offense inbounds in frontcourt.
4. Defense causes ball OOB after FG/FT miss (hits rim), offense retains possession in frontcourt.
5. AP arrow favors offense after held ball following shot that hit rim (prior to team control).
6. Offensive foul in their BACKCOURT; defense awarded ball in their frontcourt.
7. Held ball; defense awarded ball in their frontcourt.
8. Violation; defense awarded ball in their frontcourt.
9. Double personal foul, one intentional/flagrant on offense; defense awarded ball in frontcourt.
10. Shot hits rim/flange, batted into backcourt OOB by offense; defense gets ball in their frontcourt.

### RESET to 20 Seconds OR TIME REMAINING (whichever is greater)
1. Personal foul by defense, ball inbounded in frontcourt by offense (shot clock was above 20).
2. Kicked/fisted ball by defense, ball inbounded in frontcourt by offense.
3. Inadvertent whistle, no player/team possession, AP arrow favors either team in frontcourt.

### NO RESET (Unexpired Time Remains)
- Ball deflected OOB retained by offense (offense threw it OOB during a pass — stays same).
- Time-out.
- Double foul.
- Defense causes held ball during team control; AP arrow favors OFFENSE → NO reset.
- Defense commits foul/violation → frontcourt throw-in by offense, shot clock above :20 → NO reset (leave it above 20; only reset if at/below 20).

### Shot Clock Violation
- Signal: Two hands above head, index fingers extended; verbal "shot clock."
- Shot clock horn doesn't stop play — official must whistle.

### Officials Signals to Table
- Shot clock violation: two hands overhead with index fingers extended.
- Reset to 35: circle/roll hands signal (full reset).
- Reset to 20: one arm extended at 90-degree angle.

---

## 4. MECHANICS — 3-PERSON CREW POSITIONING (MSHSL)

### General Positioning
- **Lead (L):** Under the basket, baseline side opposite ball. Responsible for: end line, underneath basket, block/charge calls near basket, restricted area, post play, goaltending. Works deep — at or near the end line. Mirror ball as Lead.
- **Trail (T):** Half-court side, near 28-foot mark. Primary for: 3-point attempts (signal), perimeter fouls, out-of-bounds on sideline, bounce-pass throw-in administration.
- **Center (C):** Middle of court, opposite table. Coverage: ⅓ of court, ½ of paint area. Handles: off-ball activity, top of key, secondary coverage, post area.
- **Areas of coverage:** L = ½ lane to 3-pt line; T = above lane extended to 3-pt, end line, ⅔ behind arc; C = ⅓ court, ½ of paint.

### Key Positioning Principles
- **Lead rotation (L):** Rotate when ball crosses midline of paint and drive/post threat is imminent. If rotate late and ball goes other way (shot/turnover), rotate back. If ball reversed to far corner, Lead can rotate. When 2 posts on opposite side of floor, flex/rotate.
- **Trail (T):** Has the sideline; do not get too low. Do not get straight-lined on fast break — stop at FT line. Call out of primary only by coming in strong under the 3-pt line.
- **Center (C):** Has other sideline from Trail. Back on press defense. Count down from 5 visible at end of half. Be deliberate but not robotic.
- **Stick around huddle until 2nd horn for timeout** — get teams out on time.
- **As Lead:** If ball crosses midline of paint and drive/post threat imminent → rotate. After rotation, if ball goes away → rotate back.

### Pre-Game
- Arrive on court 15 minutes before tip.
- Space out across from table in typical pre-game positioning.
- Referee leads pregame conference with head coaches and captains; then meet with table.
- U1 meets with shot clock operator (standard protocol — review all 12 shot clock items).
- U2 observes both teams while U1 is at the table.
- Bounce pass on all throw-ins. Trail handles throw-in administration.
- Review: throw-in spots, shot clock resets, mercy rule, restricted area arc, flopping protocol.

### Free Throw Administration (3-Person)
- **Center (C):** Responsible for lane activity across AND shooter; administer at top of restricted arc; show visible count of shots; chop ball in; visible 10-second clock.
- **Lead (L):** Responsible for lane activity across; mirror ball.
- **Trail (T):** Assists with all activity; watches activity behind 3-pt line; only say "3/2/1" on last 3 FTs.
- When administering FT: do it at top of restricted arc; be loud and clear.

### Free Throw Protocols
- If FT count announced wrong, blow whistle and correct before ball becomes live.
- If free thrower loses ball while at disposal → violation.
- Visible 10-second backcourt count when shot clock is off; use shot clock when it's operating.

### Signaling & Communication
- **Full timeout:** Open hands (NOTE: fists = double foul — do NOT use fists for timeout).
- **Chop with hand closest to scorer's table** when inbounding ball (toward clock keeper).
- **Hit foul (contact foul):** Straight left arm (from reporter's view).
- **3-point attempt:** Hand up high; also for chopping, corrections, etc.
- **When at table:** Slow down; same foul call at spot and table; hands high when signaling numbers.
- **On a foul, go to table:** Show signal mirroring the play (signals that replicate play).
- Do NOT point out-of-bounds — use hand (palm toward table/floor).
- Correct your own calls (out-of-bounds reversals, etc.).
- Tick marks on floor: use for fouls in backcourt or violations in frontcourt.
- Midcourt inbound after technical foul (ball at division line opposite scorer's table).

---

## 5. RESTRICTED AREA ARC — MSHSL ADDENDUM

### Rule (MSHSL Addendum Rules 4-38, 4-41, 4-23-3)
- **Secondary defender CANNOT** establish initial legal guarding position IN the restricted area for the purpose of drawing a player-control foul/charge when defending a player with the ball (dribbling or shooting) or who has released ball for pass or try.
- If illegal contact occurs within restricted area → **BLOCKING FOUL** (except flagrant).
- **Exception 1:** Offensive player leads with foot/unnatural knee, or wards off with arm → **PLAYER-CONTROL FOUL**.
- **Exception 2:** Player in control stops continuous movement toward basket, then initiates contact with secondary defender in restricted area → **PLAYER-CONTROL FOUL**.
- **Exception 3 (Verticality):** Secondary defender in restricted area who jumps straight up with arms raised in legal vertical plane AND attempts to block a shot → VERTICALITY APPLIES. Does not apply if defender remains grounded.
- **Secondary defender definition:** (a) Teammate who helped primary defender beaten by offensive player (head/shoulders past); (b) double-teams low post player; (c) outnumbered fast-break defender (initially secondary, but may establish legal position and stay with player into arc).
- **Important:** Restricted area arc rules apply to PASS AND CRASH situations too. A secondary defender grounded in the arc trying to take a charge = blocking foul.

### Signaling Sequence for Restricted Area Block
- Option 1: Fist in air → signal block → point to restricted area (on floor below basket).
- Option 2: Fist in air → point to restricted area → signal block.
- **NOTE:** If foul is a blocking foul NOT involving the restricted area, do NOT point to the arc when signaling. That signals to partners that restricted area was a factor.
- Lead is PRIMARY on block/charge plays involving secondary defender. T and C are secondary.
- If non-calling official can provide definitive help → calling official may switch the call.

---

## 6. FLOPPING — MSHSL MECHANICS (2024-25 Adopted, Continued 2025-26)

### Rules
- **First flopping offense (by either team):** Team WARNING. Recorded in book; reported to head coach.
- **Second and subsequent offenses (same team):** TEAM TECHNICAL FOUL (2 FTs + ball at division line opposite scorer's table).
- Flopping does NOT require player to fall to floor. Includes: head bob, arm flail, dramatic reaction without contact.
- Flopping warning is NOT a violation that causes a turnover (ball does not automatically change possession on warning).
- Do NOT confuse with the five delay-of-game situations.

### Stopping Play Mechanics
- **Signal:** Show flop signal (#23/mechanic) during live play; note the time.
- **Defensive flop, offense advancing:** Do NOT stop play. Note time; issue warning at next dead ball/change of possession/when offense stops advancing.
- **Defensive flop, second offense (tech):** Stop game immediately. If shot is in air → wait for attempt to complete → penalize. Count basket if made.
- **Offensive flop, shot is in air:** Wait for rebound before issuing warning. Continuation applies.
- **Key:** Team control is lost on a shot attempt — wait for play sequence before blowing whistle.
- If warning is first offense and you stop play immediately → determine possession via AP arrow for throw-in.
- **Inbound spot for flop warning:** Tick mark where ball last was (unless shot or pass was in the air).

---

## 7. PERSONAL GAME NOTES (Selected — 2023-26 Seasons)

### 2025-26 Season

**2/13/26:** Stick around huddle until 2nd horn on timeout — ensure teams get out on time. Lead: if rotate late and ball goes away, rotate back. Pregame: start with coaches+captains, then dismiss captains for coaches-only portion.

**2/17/26:** When inbounding ball, chop with hand closest to scorer's table (clock keeper side).

**1/21/26:** Be deliberate but not robotic at the spot — use voice to explain calls/FTs/inbound spot. As Lead: if ball crosses midline of paint and drive/post threat imminent → rotate. Do NOT punch rebound foul as offensive foul — call as loose ball foul.

**12/19/25:** When running down court as new Lead, turn head and look back. Full timeout = open hands (fists = double foul). First delay of game = warning; second = technical. Example: spilling water on sideline. If other 2 officials gather, don't go over — need to watch players. Toss jump ball higher with 1 hand and more accurately.

**12/10/25:** If official administering FTs gets FT count wrong → blow whistle and correct before ball becomes live.

**West Lutheran JV/V 12/4/25:** Point at arc if restricted area blocking foul. Center: on sideline with players/not much room, move to side for better angle. Do not get too low as Trail.

**Eastview tournament 11/8/25:** Delay of game warnings, screen verbiage, contact, bench decorum, technical vs. intentional foul, post contact.

### 2024-25 Season

**If player trips on other player lying on floor = blocking foul.** Player on floor not in legal guarding position.

**Falling to ground while dribbling is NOT a travel.** Rolling on ground with ball (belly-to-back) IS a travel. Holding ball and going to ground IS travel. Dribbling and going to ground is NOT travel.

**Illegal screen:** Screener pushes defender with two hands, extends legs outside shoulder width, or extends arms making contact = team control foul.

**West Lutheran V Girls 12/17/24:** Cutting player off and reaching across body = foul. Spin dribble then 2 steps = travel. Flex/rotate as Lead when 2 posts on opposite side of floor.

**Mound Westonka 9A Boys 12/12/24:** Change hands for 5-second count to signify new count.

**Minnetonka 10A/9A Girls 12/11/24:** Only say 3/2/1 on FTs (last 3). Hit foul (contact foul) = straight left arm. Mirror ball as Lead. Number of FTs → signal toward division line when reporting. Areas: L=½ lane to 3pt; T=semicircle above lane, 3pt to endline, ⅔ behind arc; C=⅓ court, ½ paint.

**Rogers JV Girls 1/7/25:** If C: count basket and wave off game on last-second shot (if made). 1&1 handshake: don't shake — firm. Wave off shot if foul on floor.

**Secondary defender must be completely outside restricted arc on block/charge.** Pause after administering throw-in — count doesn't start right away.

**Grounded in lane = restricted area blocking foul. If jump vertically = legal.**

**If team control in frontcourt, even if tipped by other team, if offense last to touch in frontcourt and first to touch in backcourt = violation.**

**Visitation Girls V 1/23/25:** Ball reversed to far corner — Lead can rotate. Close to 10 seconds but ball pressured/trail has competitive matchup — Center needs to help.

**Edina Girls 10A/JV 1/24/25:** Call hook foul, not hold (if it's a hook).

**EP JV G 12/6/24:** Pushing player in back on layup = intentional foul.

**Fridley Girls V 1/28/25:** Remember jersey number if you call foul then tech on same player.

**Tournament 2/8/25:** Sweeping signal = player was passing. If player is pushed/swept, can use chucking signal (2 fists together, push). "Walk into" is the call — be very specific when choosing this (similarity to "walled up," coaches will question).

### Clinic Notes 2024-25

**C on 3-person FTs:** Show visible count of shots, 10-second clock, chop ball in, responsible for lane activity across AND shooter. Lead responsible for lane activity across. Trail assists all and watches behind 3-pt line.

**Tip by defender on inbounds pass is NOT a try for goal.**

**Any thrown ball from behind 3-pt arc is a 3-pt attempt** (unless deflected by offense).

**"No shot" not "on the floor" — wave off.**

**Tick marks at foul in backcourt or violation in frontcourt.**

**C back on press defense.**

**Say "no shot" when waving off.**

### Preseason/Clinic 2025-26

**Throw-in spots:** 3-pt arc instead of trapezoid — Outside (arc) = 28 ft; Inside (arc) = baseline.

**No offensive goaltending (basket interference) — 2025-26 change.**

**Slapping backboard/ring = technical foul (and basket interference if ball is in/on basket).**

**Points of emphasis:** Bench decorum and communication (pregame: only talk to head coach; warnings then techs); contact with ball handler (blow whistle early and set parameters; freedom of movement; low post contact).

---

## 8. COMMON REFEREE PITFALLS & CORRECT CALLS

1. **Fists for timeout signal** → WRONG. Open hands = full timeout. Fists = double foul.
2. **Stopping play on defensive flop while offense advancing** → WRONG. Note time; issue at dead ball.
3. **Calling offensive goaltending (2025-26)** → WRONG. Eliminated. Only defensive goaltending exists.
4. **Not pointing to restricted area when signaling restricted area block** → Missed mechanic. Signal sequence is critical.
5. **Secondary defender in restricted area jumping vertically → calling it a block** → WRONG if they jumped straight up within vertical plane attempting to block shot. Verticality applies.
6. **Treating closely guarded as applying to dribbler (MSHSL)** → WRONG. MN Mod F: 5-second closely guarded count does NOT apply while player is dribbling.
7. **Using NFHS quarter-based bonus in MSHSL games** → WRONG. MSHSL retains 1&1 on 7th team foul per HALF; double bonus on 10th.
8. **Forgetting home team wears dark in Minnesota** → WRONG vs. NFHS default. MSHSL Mod B: home = dark.
9. **Not resetting shot clock to 20 on offensive rebound that contacts rim** → WRONG. Any shot that hits rim/flange = hold button, wait for possession, then reset (off. = 20, def. = 35).
10. **Starting shot clock count immediately on jump ball** → WRONG. Shot clock starts when ball is POSSESSED on a jump ball, not when tipped.
11. **Giving throw-in to wrong team and not correcting** → Can be corrected before first dead ball / change of possession (Rule 7-6-6).
12. **Not using 3-pt line for throw-in spot (2025-26)** → Use visible 3-pt line. Inside = end line spot; outside = 28-foot mark.
13. **Calling flopping as a player technical** → WRONG since 2024-25. First offense = team WARNING; second = team technical.
14. **Home team wearing white uniforms in MSHSL** → Uniform violation. Home = dark; visitor = white (Mod B).
15. **Not having U1 pregame the shot clock operator** → Mandatory every game since 2024-25 season.

---

## 9. FRONT COURT / BACK COURT — KEY SITUATIONS

- If team in control in frontcourt: even if ball is tipped by defender, if offense is last to touch in frontcourt AND first to touch in backcourt → BACKCOURT VIOLATION.
- Ball passed from frontcourt, touches backcourt before player touches it: player who touches it in backcourt — if from the same team (offense), backcourt violation.
- Division line is part of the backcourt (Rule 4-13-2).
- Defensive player may take ball into backcourt. Backcourt violation only applies to the team that last had team control in the frontcourt.

---

## 10. COMMONLY TESTED RULE SPECIFICS — EXACT DEFINITIONS

> **HOW TO USE THIS SECTION:** These are the definitional rules that written exams
> probe and that summaries strip out. If a question matches an entry here, answer
> from it and STOP REASONING. Do not hunt for exceptions that are not listed.
> Content below is taken from the NFHS Basketball Rules Book text held in this
> project, not from recollection.

### Correctable Errors — Rule 2-10
Officials may correct an error only if a rule was inadvertently set aside and it resulted in one of these **five** situations (2-10-1):
- **a.** Failure to award a merited free throw
- **b.** Awarding an unmerited free throw
- **c.** Permitting a wrong player to attempt a free throw
- **d.** Attempting a free throw at the wrong basket
- **e.** Erroneously counting or canceling a score

**Timing limit (2-10-2):** the error must be recognized by an official **no later than during the first dead ball after the clock has properly started.**
**Special case (2-10-3):** for an erroneously counted/canceled score (1e) made while the clock was running and the ball dead, it must be recognized **before the second live ball.**

*Exam traps:* A missed foul, a wrong out-of-bounds call, or a wrong AP-arrow direction are **NOT** correctable errors — the list is exactly those five. The head coach may go to the table to request a 60-second time-out to confer about a correctable error, or to prevent/rectify a timing, scoring, or alternating-possession mistake. Assistant coaches are never authorized to approach the table.

### Airborne Shooter — Rule 4-1-1
An airborne shooter is a player who **has released the ball on a try or tap for a goal, or has tapped the ball, and has not returned to the floor.**
*Key consequence:* an airborne shooter is still considered to be in the act of shooting. A foul by an airborne shooter is a **player-control foul** (charged foul), so no free throws are awarded to the opponent unless the bonus applies for a different reason.

### Continuous Motion — Rule 4-11
- **4-11-1:** Continuous motion applies to a try or tap for field goals and free throws, but has **no significance unless a DEFENSIVE player fouls** during the interval that begins when the habitual throwing movement starts (or with the touching on a tap) and ends when the ball is clearly in flight.
- **4-11-2:** If an opponent fouls after the try/tap has started, the player may complete the customary arm movement, and if pivoting or stepping when fouled, may complete the usual foot or body movement. These privileges apply **only when the throwing motion started before the foul and before the ball is in flight.**
- **4-11-3:** Continuous motion does **NOT** apply if a **TEAMMATE** fouls after a try/tap starts and before the ball is in flight — **the ball becomes dead immediately** and the goal does not count.

*Exam trap:* the defense-vs-teammate distinction is the whole rule. Defensive foul → continuation allowed. Teammate foul → ball dead at once.

### Incidental Contact — Rule 4-27
Incidental contact is contact with an opponent which **is permitted and does not constitute a foul.**

### Foul Categories — Rule 4-19 (exact definitions)
- **4-19-2 Common foul:** a personal foul that is neither flagrant, nor intentional, nor committed against a player trying/tapping for a field goal, nor part of a double, simultaneous, or multiple foul.
- **4-19-9 False double foul:** fouls by **both** teams, the second occurring before the clock is started following the first, with at least one attribute of a double foul absent.
- **4-19-10 Simultaneous foul:** a foul by **both** teams at approximately the same time, **not** committed by opponents against each other.
- **4-19-11 Multiple foul:** **two or more teammates** commit personal fouls **against the same opponent** at approximately the same time.
- **4-19-12 False multiple foul:** two or more fouls by the **same team** where the last foul occurs before the clock is started following the first, and at least one attribute of a multiple foul is absent.
- **4-19-13 Team foul:** any personal or technical foul (**except indirect technical fouls**) charged to a team. All team fouls count toward the bonus.
- **4-19-14 Unsporting foul:** a **noncontact** technical foul consisting of unfair, unethical, or dishonorable conduct.

*Exam traps:* "Multiple" = same team, same victim. "Double" = opponents fouling each other. "Simultaneous" = both teams, but NOT against each other. Indirect technicals do **not** count as team fouls.

### Fumble — Rule 4-21
A fumble is the **accidental** loss of player control when the ball unintentionally drops or slips from a player's grasp.
*Key consequence:* a fumble is not a dribble. A player may recover their own fumble and then dribble — that is legal, and it is not a double dribble.

### Goaltending — Rule 4-22 (with 2025-26 changes layered)
- **4-22-1:** touching the ball during a try/tap while it is in **downward flight, entirely above ring level**, with a possibility of entering the basket, and not touching the cylinder.
- **4-22-2:** touching the ball **outside the cylinder** during a free-throw attempt.
- **[2025-26]** Offensive goaltending is **eliminated** — only defensive players can commit it.
- **[2025-26] 4-22-3:** once the ball contacts the **backboard**, it is automatically considered to be in downward flight.

### Coverage status of the formerly missing definitions
Section 11 below now holds the full 2026 Rules Book text for **every** definition previously identified as missing:
- **Rule 4:** 4-24 (hands/arms), 4-34 (players/bench personnel), 4-35 (player location/status), 4-36 (point of interruption), 4-37 (rebounding), 4-39 (rule), 4-40 (screen), 4-41 (shooting/try/tap), 4-42 (throw-in/thrower/designated spot), 4-43 (time-out lengths, successive time-out)
- **Rule 5:** 5-5 (length of quarter), 5-6 (beginning/ending quarter or extra period), 5-7 (extra period), 5-8 (time-out, stopping play), 5-11 (charged time-outs — allotment, durations, exceptions, excess, successive, simultaneous)
- **Rule 6:** complete — 6-1 (live ball) through 6-7 (dead ball), including jump-ball administration and alternating possession

All of the above are **CONFIRMED** sources. Cite them directly.

⚠️ **SOURCE NOTE:** an earlier version of the project's rulebook file contained an unreliable "Rule 5 (continued)" block whose numbers contradicted the 2026 Rules Book (3-minute extra periods, five 60-second time-outs — both wrong). That block has been removed from the knowledge base, but if ANY retrieved text ever conflicts with Section 11, Section 11 is transcribed from the 2026 book and is AUTHORITATIVE — use it and say so.

If a question turns on the precise wording of a rule not held here, answer **NOT IN MY KNOWLEDGE BASE** and tell the user to check the rulebook directly. Practical guidance elsewhere in this knowledge base may still be cited, but never present it as the rule's definitional text.

---

## 11. FULL DEFINITIONS — FORMERLY MISSING SECTIONS (2026 NFHS RULES BOOK)

> Transcribed verbatim from the 2026 NFHS Basketball Rules Book to fill the
> placeholder stubs identified in Section 10. Cite these as **CONFIRMED**.
> If any other retrieved source conflicts with this section, THIS SECTION IS
> AUTHORITATIVE.

## RULE 4 — DEFINITIONS

### SECTION 24 HANDS AND ARMS, LEGAL AND ILLEGAL USE (Rule 4-24)

ART. 1 . . . It is legal to extend the arms vertically above the shoulders and need not be lowered to avoid contact with an opponent when the action of the opponent causes contact. This legal use of the arms and hands usually occurs when guarding the player making a throw-in, the player with the ball in pressing tactics and a player with the ball who is maneuvering to try for goal by pivoting, jumping, etc.

ART. 2 . . . It is legal use of hands to reach to block or slap the ball controlled by a dribbler or a player throwing for goal or a player holding it and accidentally hitting the hand of the opponent when it is in contact with the ball.

ART. 3 . . . It is legal to hold the hands and arms in front of the face or body for protection and to absorb force from an imminent charge by an opponent. This same protective use of the arms and hands occurs when a player who has set a screen outside the opponent's visual field is about to be run into by the player being screened. The action, however, should be a recoil action rather than a pushing action.

ART. 4 . . . It is not legal to use hands and arms or hips and shoulders to force through a screen or to hold the screener and then to push the screener aside in order to maintain a guarding position relative to the opponent.

ART. 5 . . . It is not legal to use hands on an opponent which in any way inhibits the freedom of movement of the opponent or acts as an aid to a player in starting or stopping.

ART. 6 . . . It is not legal to extend the arms fully or partially in a position other than vertical so that the freedom of movement of an opponent is hindered when contact with the arms occurs. The extension of the elbows when the hands are on the hips or when the hands are held near the chest or when the arms are held more or less horizontally are examples of the illegal positions used.

ART. 7 . . . It is not legal to use the hand and/or forearm to prevent an opponent from attacking the ball during a dribble or when throwing for goal.

ART. 8 . . . It is not legal to swing arms and elbows excessively. This occurs when:

a. Arms and elbows are swung about while using the shoulders as pivots, and the speed of the extended arms and elbows is in excess of the rest of the body as it rotates on the hips or on the pivot foot.

b. The aggressiveness with which the arms and elbows are swung could cause injury to another player if contacted. Using this description as a basis, an official will promptly and unhesitatingly rule such action with arms and elbows a violation.

ART. 9 . . . It is not legal to lock arms or grasp a teammate(s) in an effort to restrict the movement of an opponent.

### SECTION 34 PLAYERS/BENCH PERSONNEL/SUBSTITUTES/TEAM MEMBERS (Rule 4-34)

ART. 1 . . . A player is one of five team members who are legally on the court at any given time, except during time-outs or intermissions.

ART. 2 . . . Bench personnel are all individuals who are part of or affiliated with a team, including, but not limited to: substitutes, coaches, manager(s) and statistician(s). During time-out or intermissions, all team members are bench personnel for the purpose of penalizing unsporting behavior.

ART. 3 . . . A substitute becomes a player when the substitute legally enters the court. If entry is not legal, the substitute becomes a player when the ball becomes live. A player becomes bench personnel after the substitute becomes a player or after notification of the coach following the player's disqualification.

ART. 4 . . . A team member is a member of bench personnel who is in uniform and is eligible to become a player.

### SECTION 35 PLAYER LOCATION/STATUS (Rule 4-35)

ART. 1 . . . The location/status of a player or nonplayer is determined by where the person is touching the floor as far as being:

a. Inbounds or out-of-bounds.

b. In the frontcourt or backcourt.

c. Outside (behind/beyond) or inside the three-point field-goal line.

ART. 2 . . . When a player is touching the backcourt, out of bounds or the three-point line, the player is located in backcourt, out of bounds or inside the three-point line, respectively.

ART. 3 . . . The location/status of an airborne player with reference to the three factors of Article 1 is the same as at the time such player was last in contact with the floor or an extension of the floor, such as a bleacher.

### SECTION 36 POINT OF INTERRUPTION (Rule 4-36)

ART. 1 . . . Method of resuming play due to an official's inadvertent whistle, an interrupted game, as in Rule 5-4-4, a correctable error, as in Rule 2-10-6, a double personal, double technical or simultaneous foul, as in Rules 4-19-8 and 4-19-10.

ART. 2 . . . Play shall be resumed by one of the following methods:

a. A throw-in to the team that was in control:

1. In the frontcourt at one of the four designated spots based on the ball's location relative to the three-point line when the interruption occurred. (Diagram 5)

2. In the backcourt at one of the four designated spots based on the ball's location relative to the three-point line where the ball was located when the interruption occurred.

b. A free throw or a throw-in when the interruption occurred during this activity or if a team is entitled to such.

c. A jump ball or alternating-possession throw-in when neither team is in control and no goal, infraction, nor end of quarter/extra period is involved when the game is interrupted.

ART. 3 . . . When the ball remains live after a violation or foul (as in Rule 4-19-8) during a try or tap for field goal, the point of interruption is determined to be when the ball becomes dead following the violation or foul.

### SECTION 37 REBOUNDING (Rule 4-37)

ART. 1 . . . Rebounding is an attempt by any player to secure control of the ball following a try or tap for field goal. In a rebounding situation there is no player or team control.

ART. 2 . . . To obtain or maintain legal rebounding position, a player may not:

a. Displace, charge or push an opponent.

b. Extend shoulders, hips, knees or extend the arms or elbows fully or partially in a position other than vertical so that the freedom of movement of an opponent is hindered when contact with the arms or elbows occurs.

c. Bend the body in an abnormal position to hold or displace an opponent.

d. Violate the principle of verticality.

ART. 3 . . . Every player is entitled to a spot on the playing court, provided the player gets there first without illegally contacting an opponent.

### SECTION 39 RULE (Rule 4-39)

ART. 1 . . . A rule is one of a group of regulations which governs the game.

ART. 2 . . . A game regulation, commonly called a rule, sometimes states or implies that the ball is dead or a foul or violation is involved. If it does not, it is assumed the ball is live and no foul or violation has occurred to affect the situation.

ART. 3 . . . A single infraction is not complicated by a second infraction unless so stated or implied.

### SECTION 40 SCREEN (Rule 4-40)

ART. 1 . . . A screen is legal action by a player who, while touching the playing court, without causing contact, delays or prevents an opponent from reaching a desired position.

ART. 2 . . . To establish a legal screening position:

a. The screener may face any direction.

b. Time and distance are relevant.

c. The screener shall be stationary, except when both the screener and opponent are moving in the same path and the same direction.

d. The screener shall stay within the screener's vertical plane with a stance approximately shoulder width apart.

ART. 3 . . . When screening a stationary opponent from the front or side (within the visual field), the screener may be anywhere short of contact.

ART. 4 . . . When screening a stationary opponent from behind (outside the visual field), the screener shall allow the opponent one normal step backward without contact.

ART. 5 . . . When screening a moving opponent, the screener shall allow the opponent time and distance to avoid contact by stopping or changing direction. The speed of the player to be screened will determine where the screener may take a stationary position. The position will vary and may be one to two normal steps or strides from the opponent.

ART. 6 . . . When screening an opponent who is moving in the same path and direction as the screener, the player behind is responsible if contact is made because the player in front slows up or stops and the player behind overruns the opponent.

ART. 7 . . . A player who is screened within the player's visual field is expected to avoid contact by going around the screener. In cases of screens outside the visual field, the opponent may make inadvertent contact with the screener and if the opponent is running rapidly, the contact may be severe. Such a case is to be ruled as incidental contact provided the opponent stops or attempts to stop on contact and moves around the screen, and provided the screener is not displaced if the screener has the ball.

ART. 8 . . . A player may not use the arms, hands, hips or shoulders to force through a screen or to hold the screener and then push the screener aside in order to maintain a guarding position on an opponent.

### SECTION 41 SHOOTING, TRY, TAP (Rule 4-41)

ART. 1 . . . The act of shooting begins simultaneously with the start of the try or tap for field goal and ends when the ball is clearly in flight, and includes the airborne shooter.

ART. 2 . . . A try for field goal is an attempt by a player to score two or three points by throwing the ball into a team's own basket. A player is trying for goal when the player has the ball and in the official's judgment is throwing or attempting to throw for goal. It is not essential that the ball leave the player's hand as a foul could prevent release of the ball.

ART. 3 . . . The try starts when the player begins the motion which habitually precedes the release of the ball.

ART. 4 . . . The try ends when the throw is successful, when it is certain the throw is unsuccessful, when the thrown ball touches the floor or when the ball becomes dead.

ART. 5 . . . A tap for field goal is the contacting of the ball with any part of a player's hand(s) in an attempt to direct the ball into the basket.

ART. 6 . . . A tap for field goal shall be considered the same as a try for field goal, except as in Rule 5-2-5.

ART. 7 . . . The tap for field goal starts when the player's hand(s) touches the ball.

ART. 8 . . . The tap for field goal ends in exactly the same manner as a try.

### SECTION 42 THROW-IN, THROWER, DESIGNATED SPOT (Rule 4-42)

ART. 1 . . . The thrower is the player who attempts to make a throw-in.

ART. 2 . . . A throw-in is a method of putting the ball in play from out-of-bounds.

ART. 3 . . . The throw-in and the throw-in count begin when the ball is at the disposal of a player of the team entitled to it.

ART. 4 . . . The throw-in count ends when the ball is released by the thrower so the passed ball goes directly into the court.

ART. 5 . . . The throw-in ends when:

a. The passed ball touches or is legally touched by another player inbounds.

b. The passed ball touches or is touched by another player out-of-bounds, except as in Rule 7-5-7.

c. The throw-in team commits a throw-in violation.

ART. 6 . . . The designated throw-in spot is 3 feet wide with no depth limitation and is established and signaled by the official prior to putting the ball at the thrower's disposal.

NOTE: The thrower shall keep one foot on or over the designated spot until the ball is released. The traveling and dribbling rules are not in effect for a throw-in.

### SECTION 43 TIME-OUT (Rule 4-43)

ART. 1 . . . A 60-second time-out charged to a team is a maximum of one minute in length. A 30-second time-out charged to a team is a maximum 30 seconds in length. A warning is sounded 15 seconds prior to the expiration of the 30 or 60 seconds.

ART. 2 . . . A successive time-out is one which is granted to either team before the clock has started following the previous time-out.

---

## RULE 5 — SCORING AND TIMING REGULATIONS

### SECTION 5 LENGTH OF QUARTER (Rule 5-5)

ART. 1 . . . Playing time for teams of high school age shall be four quarters of eight minutes each with intermissions of one minute after the first and third quarters, and 10 minutes between halves. The halftime intermission may be extended to a maximum of 15 minutes for special activities, provided home management has properly notified the visiting team prior to the start of the game.

ART. 2 . . . An organization sponsoring games involving teams which combine ninth-grade students with students in the eighth and/or seventh grades, may play those games in quarters of eight minutes.

ART. 3 . . . A quarter(s) may be shortened in an emergency or at any time by mutual agreement of the opposing coaches and the referee.

Note: By state association adoption, a running clock may be instituted when a specified point differential is reached at a specified time in the game.

*[MSHSL cross-reference: Minnesota varsity plays two 18-minute halves (Mod A), not quarters, and its running-clock adoption is the Mod G mercy rule — 35+ point differential with less than 9:00 remaining in the second half.]*

### SECTION 6 BEGINNING, ENDING A QUARTER OR EXTRA PERIOD (Rule 5-6)

ART. 1 . . . Each quarter or extra period begins when the ball first becomes live.

ART. 2 . . . Each quarter or extra period ends when the signal illuminates or sounds indicating time has expired, as in Rule 1-14.

EXCEPTIONS:

1. If the ball is in flight during a try or tap for field goal, the quarter or extra period ends when the try or tap ends.

2. If a held ball or violation occurs so near the expiration of time that the clock is not stopped before time expires, the quarter or extra period ends with the held ball or violation.

3. If a foul occurs so near the expiration of time that the timer cannot get the clock stopped before time expires or after time expires, but while the ball is in flight during a try or tap for field goal, the quarter or extra period ends when the free throw(s) and all related activity have been completed. No penalty or part of a penalty carries over from one quarter or extra period to the next, except when a correctable error, as in Rule 2-10, is rectified. No free throw(s) shall be attempted after time has expired for the fourth quarter or any extra period, unless the point(s) would affect the outcome of the game.

4. If a technical foul occurs after the ball becomes dead to end a quarter or extra period, the next quarter or extra period is started by administering the free throws. If the fourth quarter or extra period ends and the score is tied, the free throws will start the extra period. If the score is not tied and the results of the free throws would tie or win the game, the free throws are administered as part of the preceding quarter/period.

### SECTION 7 EXTRA PERIOD (Rule 5-7)

ART. 1 . . . If the score is tied at the end of the fourth quarter, play shall continue without change of baskets for one or more extra periods with a one-minute intermission before each extra period.

ART. 2 . . . The game ends if, at the end of any extra period, the score is not tied.

ART. 3 . . . The length of each extra period shall be four minutes (or half the time of a regulation quarter for non-varsity contests). As many such periods as are necessary to break the tie shall be played. Extra periods are an extension of the fourth quarter.

ART. 4 . . . Once the ball becomes live in the extra period, it will be played even though a correction in the fourth quarter score is made.

*[MSHSL cross-reference: the 35-second shot clock runs in extra periods (Mod E). The halves format (Mod A) does not alter Rule 5-7's extra-period provisions.]*

### SECTION 8 TIME-OUT, STOPPING PLAY (Rule 5-8)

Time-out occurs and the clock, if running, shall be stopped when an official:

ART. 1 . . . Signals:

a. A foul.

b. A held ball.

c. A violation.

d. A time-out.

ART. 2 . . . Stops play:

a. Because of an injury as in Rules 3-3-6, 3-3-7 and 3-3-8.

b. To confer with the scorer or timer.

c. Because of unusual delay in getting a dead ball live.

d. For any other situations or any emergency.

Note: When a player is injured as in Article 2(a), the official may suspend play after the ball is dead or is in control of the injured player's team or when the opponents complete a play. A play is completed when a team loses control (including throwing for goal) or withholds the ball from play by ceasing to attempt to score or advance the ball to a scoring position. When necessary to protect an injured player, the official may immediately stop play.

ART. 3 . . . Grants and signals a player's/head coach's oral or visual request for a time-out, such request being granted only when:

a. The ball is at the disposal or in control of a player of the same team.

b. The ball is dead, unless replacement of a disqualified, or injured player(s), or a player directed to leave the game is pending, and a substitute(s) is available and required.

ART. 4 . . . Responds to the scorer's signal to grant a coach's request that a correctable error, as in Rule 2-10, or a timing, scoring or alternating-possession mistake be prevented or rectified. The appeal to the official shall be presented at the scorer's table where a coach of each team may be present.

### SECTION 11 CHARGED TIME-OUTS (Rule 5-11)

ART. 1 . . . Three 60-second and two 30-second time-outs may be charged to each team during a regulation game. Each team is entitled to one additional 60-second time-out during each extra period. Unused time-outs accumulate and may be used at any time.

Note: State associations may determine the number of electronic media time-outs for games which are transmitted and may reduce the number of charged time-outs.

ART. 2 . . . A single 60-second time-out charged to a team shall not exceed one minute and shall be conducted within the confines of the time-out area. A warning signal for the teams to prepare to resume play is sounded with 15 seconds remaining. Such a time-out shall not be reduced in length unless both teams are ready to play before the time-out is over.

ART. 3 . . . A single 30-second charged time-out shall not exceed 30 seconds and players shall remain standing within the time-out area. A warning signal for teams to prepare to resume play is sounded with 15 seconds remaining. No on-court entertainment shall occur during this time.

ART. 4 . . . Only one 60-second time-out is charged (or one 30-second time-out, if that is the only type of time-out remaining) in Rule 5-8-4 regardless of the amount of time consumed when no correction is made.

EXCEPTION: No time-out is charged:

a. If, in Rule 5-8-3, the player's request results from displaced eyeglasses or lens.

b. If, in Rule 5-8-4, the error or mistake is prevented or rectified.

ART. 5 . . . A time-out shall not be granted until after the ball has become live to start the game. The additional 60-second time-out provided for each extra period(s) shall not be granted until after the ball has become live to start the extra period(s).

ART. 6 . . . Time-outs in excess of the allotted number may be requested and shall be granted during regulation playing time or any extra period at the expense of a technical foul for each, as in Rule 10-2-3.

ART. 7 . . . Successive time-outs, as in Rule 4-43-2, shall not be granted after the expiration of playing time for the fourth quarter or any extra period. In all other instances, they shall be administered in the order in which they were requested.

ART. 8 . . . Time-outs simultaneously requested by opposing teams or those requested to keep players in the game that were directed to leave for injury/blood, as in Rules 3-3-6 and 3-3-7 Notes, shall be granted, charged to the respective team and administered concurrently. When one team is charged with a 30-second time-out and the other a 60-second time-out, the duration shall be 60 seconds.

*[MSHSL cross-reference: the 5-11-1 allotment (3×60 + 2×30, +1×60 per extra period, unused accumulate) applies in Minnesota. Note 5-8-3b — a time-out may be granted on a dead ball only when no required substitution is pending — and 5-11-6, where a time-out beyond the allotment is granted but costs a technical foul (10-2-3).]*

---

## RULE 6 — LIVE BALL AND DEAD BALL

### SECTION 1 LIVE BALL (Rule 6-1)

ART. 1 . . . The game and each extra period shall be started by a jump ball in the center restraining circle. After any subsequent dead ball, the only way to get the ball live is to resume play by a jump ball in the center restraining circle, by a throw-in or by a free throw. The dribble and traveling rules are not in effect in these situations.

ART. 2 . . . The ball becomes live when:

a. On a jump ball, the tossed ball leaves the official's hand(s).

b. On a throw-in, it is at the disposal of the thrower.

c. On a free throw, it is at the disposal of the free thrower.

Note: Any rules statement is made on the assumption that no infraction is involved unless mentioned or implied. If such infraction occurs, the rule governing it is followed. For example, a game or extra period will not start with a jump ball if a foul occurs before the ball becomes live.

### SECTION 2 STARTING GAME/QUARTER/EXTRA PERIOD (Rule 6-2)

ART. 1 . . . The game, quarter and each extra period begins when the ball becomes live as specified in Rule 6-1-2 for a jump ball, throw-in or free throw.

ART. 2 . . . To start the game and each extra period, the ball shall be put in play in the center restraining circle by a jump ball between any two opponents.

ART. 3 . . . To start the second, third and fourth quarters, the ball shall be put in play by a throw-in under the alternating-possession procedure.

### SECTION 3 JUMP-BALL ADMINISTRATION (Rule 6-3)

ART. 1 . . . For any jump ball, each jumper shall have both feet within that half of the center restraining circle which is farther from the jumper's basket.

ART. 2 . . . When the official is ready and until the ball is tossed, nonjumpers shall not:

a. Move onto the center restraining circle (within 3 feet).

b. Change position around the center restraining circle.

ART. 3 . . . Teammates may not occupy adjacent positions around the center restraining circle if an opponent indicates a desire for one of these positions before the official is ready to toss the ball.

ART. 4 . . . The ball shall be tossed upward between the jumpers in a plane at right angles to the sidelines. The toss shall be to a height greater than either of them can jump so that it will drop between them.

ART. 5 . . . Until the tossed ball is touched by one or both jumpers, nonjumpers shall not:

a. Have either foot break the plane of the center restraining circle cylinder.

b. Take a position in any occupied space (within 3 feet of the center restraining circle).

ART. 6 . . . The tossed ball shall be touched by one or both of the jumpers after it reaches its highest point. If the ball contacts the floor without being touched by at least one of the jumpers, the official shall toss it again.

ART. 7 . . . Neither jumper shall:

a. Touch the tossed ball before it reaches its highest point.

b. Leave the center restraining circle until the ball has been touched.

c. Catch the ball before the jump ball ends.

d. Touch the ball more than twice.

ART. 8 . . . The jump ball and the restrictions in Rule 6-3-7 end when the touched ball contacts one of the eight nonjumpers, an official or the floor.

Note: During a jump ball, a jumper is not required to face the jumper's own basket, provided the jumper is in the proper half of the center restraining circle. The jumper is also not required to jump and attempt to touch the tossed ball. However, if neither jumper touches the ball it should be tossed again with both jumpers being ordered to jump and try to touch the ball.

### SECTION 4 ALTERNATING POSSESSION (Rule 6-4)

ART. 1 . . . Other than the start of the game and each extra period, the teams will alternate taking the ball out of bounds for a throw-in. The team obtaining control from the jump ball establishes the alternating-possession procedure, and the arrow is set toward the opponent's basket. Control may also be established by the results of a violation or foul, as in Rule 4-3.

ART. 2 . . . To start the second, third and fourth quarters, the throw-in shall be from out of bounds at the division line opposite the scorer's and timer's table.

ART. 3 . . . Alternating-possession throw-ins shall be from:

a. One of the four designated spots based on the ball's location relative to the three-point line if team control is retained or gained in the team's frontcourt. (Diagram 5)

b. The out-of-bounds spot nearest to where the ball was located if team control is retained or gained in the team's backcourt.

ART. 4 . . . An alternating-possession throw-in shall result when:

a. A held ball occurs.

b. The ball goes out of bounds, as in Rule 7-3.

c. Simultaneous floor or free-throw violations occur.

d. A live ball lodges between the backboard and ring or comes to rest on the flange, unless a free throw or throw-in follows.

e. Opponents commit simultaneous basket-interference violations.

f. The point of interruption cannot be determined as in Rule 4-36-2c.

g. A warning is issued for faking being fouled when an offensive player fakes being fouled after a try has been released and the attempt is unsuccessful.

Note: When the alternating-possession procedure has not been established, the jump ball shall be in the center restraining circle between the two players involved in the previous action.

ART. 5 . . . The direction of the possession arrow is reversed immediately after an alternating-possession throw-in ends. An alternating-possession throw-in ends when the throw-in ends, as in Rule 4-42-5.

ART. 6 . . . The opportunity to make an alternating-possession throw-in is lost if the throw-in team violates. If either team fouls during an alternating-possession throw-in, it does not cause the throw-in team to lose the possession arrow. If the defensive team commits a violation during the throw-in, the possession arrow is not switched.

### SECTION 5 BALL IN PLAY BY THROW-IN (Rule 6-5)

The ball shall be put in play by a throw-in under circumstances as outlined in Rules 6-4-1, 6-4-2, 6-4-4; Rules 7-1 through 7-6; Rule 8-5; and Rules 9-1 through 9-13.

### SECTION 6 BALL IN PLAY BY FREE THROW (Rule 6-6)

The ball shall be put in play by placing it at the disposal of the free thrower before each free throw.

### SECTION 7 DEAD BALL (Rule 6-7)

The ball becomes dead, or remains dead, when:

ART. 1 . . . A goal, as in Rule 5-1, is made.

ART. 2 . . . It is apparent the free throw will not be successful on a:

a. Free throw which is to be followed by another free throw.

b. Free throw which is to be followed by a throw-in.

ART. 3 . . . A held ball occurs, or the ball lodges between the backboard and ring or comes to rest on the flange.

ART. 4 . . . A player-control or team-control foul occurs.

ART. 5 . . . An official's whistle is blown (see exceptions a and b below).

ART. 6 . . . Time expires for a quarter or extra period (see EXCEPTION a below).

ART. 7 . . . A foul, other than player-control or team-control, occurs (see exceptions a, b and c below).

ART. 8 . . . A free-throw violation by the throwing team, as in Rule 9-1, occurs.

ART. 9 . . . A violation, as in Rules 9-2 through 13, occurs (see EXCEPTION d below).

EXCEPTION: The ball does not become dead until the try or tap for field goal ends, or until the airborne shooter returns to the floor, when:

a. Article 5, 6, or 7 occurs while a try or tap for a field goal is in flight.

b. Article 5 or 7 occurs while a try for a free throw is in flight.

c. Article 7 occurs by any opponent of a player who has started a try or tap for field goal (is in the act of shooting) before the foul occurred, provided time did not expire before the ball was in flight. The trying motion shall be continuous and begins after the ball comes to rest in the player's hand(s) on a try or touches the hand(s) on a tap, and is completed when the ball is clearly in flight. The trying motion may include arm, foot or body movements used by the player when throwing the ball at the player's basket.

d. Article 9 as in Rules 9-3-3 or 9-13-1, occurs by an opponent.

Note: If A1's try or tap for field goal is legally touched in flight, the goal counts if made, if the period/quarter ends before or after the legal touching. If the touching is interference or goaltending by Team A, no points are scored. If Team B violates, the points are awarded — either two or three depending on whether it was a two or three-point try or tap for field goal.

---

## SYSTEM PROMPT INSTRUCTIONS FOR REFBUDDY

You are RefBuddy, a hyper-precise Minnesota high school basketball referee assistant.

CRITICAL LAYERING RULE: The CORE_KNOWLEDGE contains a `2023-2024_NFHS_Basketball_Rulebook.md` baseline (Sections 1–9) plus 2023-2026 changes (Section 0) at the top.
- DEFAULT to the `2023-2024_NFHS_Basketball_Rulebook.md` for any rule not listed in Section 0.
- If Section 0 contains a change for that rule, APPLY the updated rule and cite the year: e.g., "[2025-26 change]" or "[2025-26 change — overrides 2023-24 Rule X-X-X]".
- MSHSL Minnesota Modifications ALWAYS take precedence over NFHS rules for MSHSL games.
- Never cite NFHS language when MSHSL has a conflicting modification.

Your behavior:
1. Start EVERY response with the most relevant rule citation (e.g., "Rule 4-22-1 [2025-26 change]" or "MSHSL Mod G") and include the year if the rule changed.
2. Reference personal game notes when applicable.
3. End EVERY response with: "*Not official MSHSL interpretation — confirm with your assignor.*"
4. Temperature = 0 mindset: maximum precision, no guessing, no hallucinating.
5. If game context (quarter vs. half, crew size, level, MSHSL vs. NFHS) is missing, ask before ruling.
6. For video/film analysis: always include a VISIBILITY CHECK section. Use "Frame N" format.
7. For RefGrade evaluations: structured scores (0-100), frame-by-frame highlights, visibility notes, "What to work on" bullets.
8. MSHSL CRITICAL DIFFERENCES TO ALWAYS REMEMBER:
   - Home = dark; Visitor = white (Mod B)
   - Game = two 18-min halves (Mod A)
   - Bonus = 1&1 on 7th/half; 2-shot on 10th/half (Mod K) NOT quarter-based
   - Coaching box = 14 feet (Mod I)
   - No closely guarded on dribbler (Mod F)
   - Shot clock 35 seconds mandatory all varsity (Mod E)
   - Mercy rule: 35-pt lead, <9 min remaining = running clock (Mod G)
   - Restricted area arc required all levels (Mod H)
   - Apparel must be black or white (Mod C)
"""

# =============================================================================
# PROMPT CACHING
#
# CORE_KNOWLEDGE is byte-identical on every API call the app makes. Sending it
# as a cached block means Anthropic stores the tokenized prefix and charges a
# fraction of the normal input rate on subsequent hits, instead of re-billing
# ~9,000 tokens every single time. On most models cached reads also don't count
# against the input-tokens-per-minute limit.
#
# Order matters: the cached block MUST come first so the prefix match hits even
# though the task-specific instructions that follow differ per call.
# =============================================================================

CACHED_KB_BLOCK = {
    "type": "text",
    "text": CORE_KNOWLEDGE,
    "cache_control": {"type": "ephemeral"},
}


def system_blocks(instructions: str) -> list:
    """
    Build the `system` parameter for any API call.

    Block 1: CORE_KNOWLEDGE, cached (identical across every call in the app)
    Block 2: the task-specific instructions, not cached (small and varied)
    """
    return [
        CACHED_KB_BLOCK,
        {"type": "text", "text": instructions},
    ]


# =============================================================================
# SYSTEM PROMPTS
# =============================================================================

SYSTEM_PROMPT = f"""You are RefBuddy — a straightforward, hyper-precise Minnesota high school basketball referee assistant.

════════════════════════════════════════════════════════════════════════
ACCURACY GUARDRAIL — THIS OVERRIDES EVERY OTHER INSTRUCTION BELOW
════════════════════════════════════════════════════════════════════════

A confident wrong answer is far worse than "I don't know." Officials act on
what you say. These rules are absolute:

1. DO NOT INVENT EXCEPTIONS.
   If a rule states "always," "never," or "shall," and the CORE_KNOWLEDGE
   lists no exception, then THERE IS NO EXCEPTION. Do not reason your way to
   one from geometry, physics, edge cases, or what "would make sense." If you
   find yourself constructing a scenario the knowledge base never mentions in
   order to justify an answer, STOP — that is the signal you are fabricating.

2. DO NOT FABRICATE RULE NUMBERS.
   Only cite a rule number that literally appears in CORE_KNOWLEDGE. Never
   pair a real rule number with content it does not contain. If you know the
   substance but not the citation, say so instead of inventing one.

3. START EVERY RULES ANSWER WITH A CONFIDENCE PREFIX:
   • **CONFIRMED** — the rule and its exact language are in CORE_KNOWLEDGE.
   • **LIKELY** — the principle is in CORE_KNOWLEDGE but the exact wording,
     numbering, or a specific sub-case is not. Say what you are unsure of.
   • **NOT IN MY KNOWLEDGE BASE** — say this plainly, do not guess, and tell
     the user to check the current NFHS Rules Book or ask their assignor.
   Section 11 holds full 2026 text for the formerly missing definitions —
   cite those as CONFIRMED. Section 10 lists what is still NOT held; if the
   question turns on one of those, the answer is NOT IN MY KNOWLEDGE BASE.

4. TRUE/FALSE AND MULTIPLE-CHOICE QUESTIONS.
   Exam statements are frequently near-verbatim rule text, in which case the
   answer is simply TRUE. The presence of "always" or "never" is NOT evidence
   that a statement is false — many rules genuinely are absolute. Do not treat
   a T/F question as a puzzle with a hidden trick. Match it against
   CORE_KNOWLEDGE; if it matches, answer accordingly; if the rule is not
   there, say NOT IN MY KNOWLEDGE BASE rather than reasoning toward a guess.

5. DO NOT CONFUSE ADJACENT CONCEPTS.
   Game clock vs. shot clock. Player control vs. team control. Multiple vs.
   double vs. simultaneous foul. Defensive vs. teammate foul under continuous
   motion. Player vs. bench personnel. Getting these backwards is the most
   common way a fluent answer becomes a wrong one.

6. MATCH FORMAT TO CERTAINTY.
   Do not dress a guess in headings, comparison tables, and confident section
   titles. Elaborate structure implies knowledge you may not have. When
   uncertain, answer briefly and say why you are uncertain.

════════════════════════════════════════════════════════════════════════

You ONLY reference information from the CORE_KNOWLEDGE below. Cite page/rule number every time.
For video questions, first ask for transcription or key timestamps.
Never hallucinate MSHSL or NFHS mechanics.
Always ask clarifying questions on game context before ruling.

CRITICAL LAYERING RULE: The CORE_KNOWLEDGE contains a `2023-2024_NFHS_Basketball_Rulebook.md` baseline (Sections 1–9) plus a 2023–2026 changes section (Section 0) at the top.
- DEFAULT to the `2023-2024_NFHS_Basketball_Rulebook.md` for any rule not listed in Section 0.
- If Section 0 contains a change for that rule, APPLY the updated rule and cite the year: e.g., "[2025-26 change]" or "[2025-26 change — overrides Rule X-X-X]".
- MSHSL Minnesota Modifications ALWAYS override NFHS defaults. Always apply MSHSL mods for MSHSL games.
- If a rule was changed multiple times, apply the MOST RECENT version and note the history.

Your behavior:
1. Start EVERY response with the most relevant rule citation (e.g., "Rule 9-7 [Three Seconds]" or "MSHSL Mod K" or "MSHSL Shot Clock Protocol — Reset to 20") and include the year if the rule changed after 2023-24.
2. Reference personal game notes when applicable.
3. End EVERY response with: "*Not official MSHSL interpretation — confirm with your assignor.*"
4. Temperature = 0 mindset: maximum precision, no guessing, no hallucinating.
5. If game context (level, MSHSL vs. NFHS, crew size, quarter/half) is missing, ask before ruling.
6. For video/film analysis: always include a VISIBILITY CHECK section. Use "Frame N" format.
7. For RefGrade evaluations: structured scores (0-100), frame-by-frame highlights, visibility notes, "What to work on" bullets.

"""

REFGRADE_PROMPT = f"""You are RefBuddy acting as a professional officiating evaluator for Minnesota high school basketball.

Output EXACTLY this structure:

## 📊 RefGrade Report
**Clip:** [filename] | **Evaluated:** [scope] | **Frames:** [range] | **Date:** [today]

## 👁️ Visibility Check
List each position: CLEARLY VISIBLE (frames N...) / PARTIALLY VISIBLE (frames N-N) / NOT VISIBLE IN ANY FRAME

## 📈 Scores
| Category | Score | Notes |
|----------|-------|-------|
| Positioning | XX/100 | |
| Call Accuracy | XX/100 | |
| Mechanics Execution | XX/100 | |
| Dead-ball Officiating | XX/100 | |
| Communication/Signals | XX/100 | |
| **Overall** | **XX/100** | |

90-100=Excellent; 80-89=Good; 70-79=Average; 60-69=Needs work; <60=Significant concern

## 🎬 Frame-by-Frame Highlights
## ✅ Strengths
## 🔧 What to Work On
## 📋 Summary

*Not official MSHSL interpretation — confirm with your assignor.*

Cite NFHS rules and MSHSL mechanics on every observation. Never hallucinate.
"""

QUIZ_SYSTEM_PROMPT = f"""You are RefBuddy Quiz Engine — a precise question generator for Minnesota high school basketball officials.

ABSOLUTE RULES — violating these will cause test failures:
0. ACCURACY OVERRIDES EVERYTHING. Only write questions whose answers are
   verifiable in CORE_KNOWLEDGE. Never invent a rule number, never invent an
   exception to a rule stated as absolute, and never write a question about a
   definition listed in Section 10 as NOT held in the knowledge base. If you
   cannot verify the answer from CORE_KNOWLEDGE, write a different question.
   The rule_citation field must contain only citations that literally appear
   in CORE_KNOWLEDGE.
1. Respond with ONLY valid JSON. Zero preamble. Zero markdown fences. Zero trailing text.
2. Multiple-choice: EXACTLY 4 options (A, B, C, D). Exactly ONE correct answer.
3. True/False: EXACTLY 2 options: {{"A": "True", "B": "False"}}.
4. Mix types roughly 50% multiple_choice / 50% true_false. Vary the ratio naturally.
5. Questions must be CHALLENGING — not trivial. Use specific rule numbers, shot clock resets, timing rules, and realistic scenario language.
6. NEVER repeat the same topic, scenario, or rule in a batch. Cover wide breadth.
7. Distractors for MC must be plausible but clearly wrong to someone who studied.

Single question JSON structure:
{{
  "question": "Full question text — be specific and scenario-based when possible",
  "type": "multiple_choice",
  "options": {{"A": "option", "B": "option", "C": "option", "D": "option"}},
  "correct": "B",
  "explanation": "Thorough explanation: why correct answer is right, why each wrong answer is wrong, what the rule actually says.",
  "rule_citation": "Exact rule number or MSHSL Modification letter",
  "personal_note": "The specific situation or interpretation from the knowledge base, stated directly with no prefix — do NOT begin with 'From your notes' or 'From your 2024-25 notes' etc., because the UI already prepends that label (empty string if not applicable)",
  "topic": "Rules|Mechanics|Shot Clock|Positioning|Signals|Game Situations|MSHSL Specific|2025-26 Changes"
}}

True/False structure (type must be "true_false"):
{{
  "question": "True or False: [specific statement that requires knowledge to evaluate]",
  "type": "true_false",
  "options": {{"A": "True", "B": "False"}},
  "correct": "A",
  "explanation": "...",
  "rule_citation": "...",
  "personal_note": "",
  "topic": "Rules"
}}

For a BATCH of 10 questions: JSON array of 10 objects. Include:
- At least 2 MSHSL-specific questions (mercy rule, MN modifications, shot clock)
- At least 2 mechanics/positioning questions (Lead/Trail/Center responsibilities)
- At least 1 question on 2025-26 rule changes
- At least 1 scenario-based game situation question
- At least 1 question from personal game notes
- At least 1 shot clock reset scenario question
- The rest from NFHS rules (varied — not all from Rule 9)

"""

CREW_EVAL_PROMPT = f"""You are RefBuddy acting as a professional officiating evaluator for Minnesota high school basketball. You are analyzing game film to evaluate the officiating crew.

Generate a comprehensive crew evaluation report with the following structure:

## 📊 Crew Evaluation Report
**Game Film:** [filename] | **Date:** [today] | **Evaluated By:** RefBuddy

## 👁️ Visibility Check
For each crew position, note: CLEARLY VISIBLE (frames N...) / PARTIALLY VISIBLE / NOT VISIBLE — analysis inferred from play action

## 📈 Overall Crew Score: XX/100

## 📋 Per-Position Highlights
For each visible official (Lead/Trail/Center): what they did well, positioning observations, any missed calls or mechanics issues. Cite specific frames and NFHS/MSHSL rules.

## 🎬 Key Play Analysis
Walk through 3-5 significant plays/moments from the film with specific frame citations, what happened, what the correct mechanics called for, and how the officials responded.

## ✅ Crew Strengths

## 🔧 Areas for Development
Actionable bullets with specific mechanic/rule citations and suggested focus for next game.

## 📋 Summary

---
*Not official MSHSL interpretation — confirm with your MSHSL district assignor.*

Cite NFHS rules and MSHSL mechanics on every observation. Never hallucinate. Use "Frame N" format throughout.
"""

REF_EVAL_PROMPT = f"""You are RefBuddy acting as a professional officiating evaluator for Minnesota high school basketball. You are analyzing game film to evaluate ONE specific official.

Generate a focused evaluation report with the following structure:

## 📊 Official Evaluation Report
**Game Film:** [filename] | **Position Evaluated:** [position] | **Date:** [today]

## 👁️ Visibility Check
How clearly is this official visible in the provided frames? List specific frames where they appear.

## 📈 Position Score: XX/100

## 📐 Positioning Analysis
Was the official in the correct position for each situation? Cite specific frames. Reference MSHSL mechanics manual standards for this position (Lead/Trail/Center).

## 📋 Call Accuracy
Any whistles blown or situations where a whistle should have been blown. Was each decision correct per NFHS rules and MSHSL modifications? Cite Rule numbers and MSHSL Mods.

## ⚙️ Mechanics Execution
Signals, whistle timing, throw-in administration, FT administration, relay mechanics, communication. What was correct? What needs work?

## ✅ Strengths

## 🔧 Development Points
Specific, actionable improvements with exact mechanic citations and suggested drills.

## 📋 Summary

---
*Not official MSHSL interpretation — confirm with your MSHSL district assignor.*

Cite NFHS rules and MSHSL mechanics specifically. Use "Frame N" format. Never hallucinate.
"""

PREGAME_MEETING_PROMPT = f"""You are RefBuddy acting as a Minnesota high school basketball officiating coordinator.
Generate a CONCISE pre-game crew meeting agenda — maximum 1 to 1.5 printed pages.
Short bullet points only. No paragraphs. No explanations. No filler.
Each bullet must be actionable and specific. Cite rule numbers inline (e.g. MSHSL Mod G, Rule 9-7).
Total output should be ~300-400 words maximum.

Output EXACTLY this structure:

---
# Pre-Game Meeting Agenda
{{date}} | {{crew}} | {{level}}

## 2025-26 Rule Changes (know these cold)
- [list only the 2-3 most important 2025-26 changes with rule #]

## Key Mechanics Reminders
- [4-6 bullet points covering the highest-leverage mechanics for this crew size — Lead/Trail/Center specific]

## MSHSL Modifications to Confirm
- [3-4 critical MN mods that differ from NFHS — bonus system, home/visitor colors, closely guarded, mercy rule]

## Shot Clock Pre-Game
- [3-4 shot clock reminders — reset triggers, operator check, horn ≠ violation]

## Watch-Fors Tonight
- [3-5 specific situations from CORE_KNOWLEDGE most likely to come up]

## Assignor Notes
[ASSIGNOR_NOTES_PLACEHOLDER]

## Quick Scenarios (discuss briefly if time)
- [2 short scenario questions, one sentence each]

---
*Not official MSHSL interpretation.*

Keep it tight. Referees are reading this standing on a sideline before tip.
"""


# =============================================================================
# PAGE CONFIG
# =============================================================================

st.set_page_config(
    page_title="RefBuddy — MN HS Basketball",
    page_icon="🏀",
    layout="wide",
    initial_sidebar_state="expanded",
)


# =============================================================================
# CSS — v1.0 Basketball Theme (same design system as Football v3.1)
# =============================================================================

BLUE   = "#003087"   # Deep navy
BLUE_L = "#1E56A0"
CREAM  = "#FAFAF7"
CARD   = "#FFFFFF"
BORDER = "#DDE3F0"
TEXT   = "#1F2937"
MUTED  = "#4B5563"
GREEN  = "#15803D"
AMBER  = "#92400E"
RED    = "#991B1B"

# Basketball court-inspired subtle SVG background
_SVG = (
    "<svg xmlns='http://www.w3.org/2000/svg' width='120' height='120'>"
    "<rect width='120' height='120' fill='none'/>"
    # Court lines
    "<line x1='0' y1='60' x2='120' y2='60' stroke='%23003087' stroke-width='0.4' opacity='0.07'/>"
    "<line x1='60' y1='0' x2='60' y2='120' stroke='%23003087' stroke-width='0.4' opacity='0.07'/>"
    # Lane lines suggestion
    "<rect x='44' y='30' width='32' height='60' fill='none' stroke='%23003087' stroke-width='0.3' opacity='0.04'/>"
    # Half-court circle
    "<circle cx='60' cy='60' r='15' fill='none' stroke='%23003087' stroke-width='0.35' opacity='0.05'/>"
    # Basketball
    "<circle cx='60' cy='60' r='7' fill='none' stroke='%23003087' stroke-width='0.4' opacity='0.06'/>"
    "</svg>"
)
BG_URL = "data:image/svg+xml," + urllib.parse.quote(_SVG)

# ── Layer 1: Mandatory button + selectbox + sidebar + dark text ──────────────
st.markdown("""
<style>
    /* Light bg, black border, black text on ALL buttons */
    .stButton button, button, .stButton>button {
        color: #1F2937 !important;
        background-color: #F8FAFC !important;
        border: 2px solid #1F2937 !important;
        font-weight: 600;
    }
    .stButton button:hover { background-color: #E2E8F0 !important; }
    .stButton button:disabled, .stButton>button:disabled {
        background-color: #F1F5F9 !important;
        color: #94A3B8 !important;
        border-color: #94A3B8 !important;
    }
    /* PRIMARY buttons — solid navy. Used for the main call-to-action on each
       screen (Analyze, Run RefGrade, Generate) and to mark the active sub-tab
       in Ref Hub. Declared after the base rule so it wins. */
    .stButton button[kind="primary"],
    .stButton button[data-testid="stBaseButton-primary"],
    [data-testid="stBaseButton-primary"] {
        background-color: #003087 !important;
        color: #FFFFFF !important;
        border: 2px solid #003087 !important;
        font-weight: 700 !important;
        letter-spacing: 0.01em;
        box-shadow: 0 2px 8px rgba(0,48,135,0.25) !important;
    }
    .stButton button[kind="primary"]:hover,
    [data-testid="stBaseButton-primary"]:hover {
        background-color: #002266 !important;
        border-color: #002266 !important;
        color: #FFFFFF !important;
        box-shadow: 0 4px 14px rgba(0,48,135,0.35) !important;
    }
    .stButton button[kind="primary"] p,
    .stButton button[kind="primary"] div,
    [data-testid="stBaseButton-primary"] p,
    [data-testid="stBaseButton-primary"] div {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    .stButton button[kind="primary"]:disabled,
    [data-testid="stBaseButton-primary"]:disabled {
        background-color: #CBD5E1 !important;
        border-color: #CBD5E1 !important;
        color: #64748B !important;
        box-shadow: none !important;
    }
    /* Selectbox + multiselect — light bg, black border, black text */
    .stSelectbox > div, .stMultiSelect > div,
    .stSelectbox > div > div, .stMultiSelect > div > div {
        color: #1F2937 !important;
        background-color: #F8FAFC !important;
        border: 2px solid #1F2937 !important;
    }
    .stSelectbox label, .stMultiSelect label,
    [data-baseweb="select"] span, [data-baseweb="select"] div,
    [data-baseweb="popover"] li, [data-baseweb="menu"] li {
        color: #1F2937 !important;
        background-color: #F8FAFC !important;
    }
    /* Sidebar */
    [data-testid="stSidebar"] { background-color: #FFFFFF !important; }
    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] .stSelectbox,
    [data-testid="stSidebar"] .stTextInput,
    [data-testid="stSidebar"] .stMarkdown h1,
    [data-testid="stSidebar"] .stMarkdown h2,
    [data-testid="stSidebar"] .stMarkdown p,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] div { color: #1F2937 !important; }
    [data-testid="stSidebar"] [data-baseweb="select"] > div,
    [data-testid="stSidebar"] [data-baseweb="select"] span {
        color: #1F2937 !important; background-color: #F8FAFC !important;
    }
    /* Tab labels */
    .stTabs [data-baseweb="tab"] { color: #1F2937 !important; }
    .stTabs [aria-selected="true"] {
        color: #003087 !important; border-bottom: 3px solid #003087 !important;
    }
    /* Dark text everywhere */
    .stMarkdown, .stMarkdown p, .stMarkdown li, .stMarkdown h1,
    .stMarkdown h2, .stMarkdown h3, .stMarkdown h4,
    .stMarkdown span, .stMarkdown strong, .stMarkdown em,
    p, span, label, h1, h2, h3, h4, h5,
    div[data-testid="stMarkdownContainer"] p,
    div[data-testid="stMarkdownContainer"] li,
    div[data-testid="stMarkdownContainer"] span,
    .stChatMessage p, .stChatMessage span, .stChatMessage li,
    [data-testid="stChatMessageContent"] p,
    [data-testid="stChatMessageContent"] span,
    [data-testid="stChatMessageContent"] li { color: #1F2937 !important; }
</style>
""", unsafe_allow_html=True)

# ── Layer 1b: Radio, inputs, sidebar, alerts ──────────────────────────────────
st.markdown("""
<style>
/* Radio labels */
.stRadio label, .stRadio label span, .stRadio label p,
div[data-testid="stRadio"] label span,
div[data-testid="stRadio"] label p,
.stRadio > div > label > div > p {
    color: #1F2937 !important;
    font-size: 0.95rem !important;
}
/* Chat input */
.stChatInput textarea, .stChatInput input {
    color: #1F2937 !important; background-color: #FFFFFF !important;
}
/* Text areas / inputs */
.stTextArea textarea, .stTextInput input {
    color: #1F2937 !important; background-color: #FFFFFF !important;
}
/* Select boxes */
[data-baseweb="select"] span, [data-baseweb="select"] div { color: #1F2937 !important; }
/* Sidebar */
[data-testid="stSidebar"] p, [data-testid="stSidebar"] span,
[data-testid="stSidebar"] label, [data-testid="stSidebar"] div { color: #1F2937 !important; }
/* Alert boxes */
.stAlert p, .stAlert span, .stAlert div { color: #1F2937 !important; }
/* Expander headers */
.streamlit-expanderHeader p, .streamlit-expanderHeader span { color: #003087 !important; }
/* Caption */
.stCaption, .stCaption p { color: #4B5563 !important; }
</style>
""", unsafe_allow_html=True)

# ── Layer 2: Full theme CSS ───────────────────────────────────────────────────
st.markdown(f"""
<style>
html, body, [data-testid="stAppViewContainer"] {{
    background-color: {CREAM};
    background-image: url("{BG_URL}");
    background-repeat: repeat;
    color: {TEXT};
    font-family: 'Inter', 'Segoe UI', sans-serif;
}}
.main .block-container {{ background: transparent; padding-top: 0.5rem; max-width: 1100px; }}

/* Sidebar */
[data-testid="stSidebar"] {{
    background-color: {CARD}; border-right: 2px solid {BORDER};
    box-shadow: 2px 0 8px rgba(0,48,135,0.06);
}}

/* Hero */
.home-hero {{ text-align: center; padding: 2.2rem 2rem 1.4rem 2rem; }}
.home-hero-title {{
    color: {BLUE} !important; font-size: 3.2rem; font-weight: 900;
    letter-spacing: -1.5px; margin: 0 0 0.2rem 0; line-height: 1.1;
}}
.home-hero-slogan {{ color: {MUTED}; font-size: 1.1rem; font-weight: 500; margin: 0 0 1.6rem 0; }}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {{
    background-color: {CARD}; border-bottom: 2px solid {BORDER};
    border-radius: 8px 8px 0 0; gap: 2px; padding: 0 0.4rem;
}}
.stTabs [data-baseweb="tab"] {{
    color: {MUTED} !important; font-weight: 600; font-size: 0.9rem;
    padding: 0.55rem 1rem; border-radius: 6px 6px 0 0;
}}
.stTabs [aria-selected="true"] {{
    color: {BLUE} !important; background-color: {CREAM} !important;
    border-bottom: 3px solid {BLUE} !important;
}}

/* Cards */
.rb-card {{
    background: {CARD}; border: 1px solid {BORDER}; border-radius: 10px;
    padding: 1.2rem 1.4rem; margin-bottom: 1rem;
    box-shadow: 0 2px 8px rgba(0,48,135,0.05); color: {TEXT};
}}
.rb-card-blue {{
    background: {CARD}; border-left: 4px solid {BLUE};
    border-radius: 0 10px 10px 0; padding: 1rem 1.2rem; margin-bottom: 0.8rem;
    box-shadow: 0 2px 8px rgba(0,48,135,0.05); color: {TEXT};
}}

/* Report output */
.report-output {{
    background: {CARD}; border: 2px solid {BORDER}; border-radius: 10px;
    padding: 1.6rem 2rem; margin-top: 1rem;
    box-shadow: 0 3px 12px rgba(0,48,135,0.07); color: {TEXT};
    line-height: 1.7; font-size: 0.93rem;
}}
.report-output h1, .report-output h2, .report-output h3, .report-output h4 {{
    color: {BLUE} !important;
}}

/* Quiz cards */
.quiz-question-card {{
    background: {CARD}; border: 2px solid {BORDER}; border-radius: 12px;
    padding: 1.5rem 1.8rem; margin-bottom: 1rem;
    box-shadow: 0 3px 12px rgba(0,48,135,0.08); color: {TEXT};
}}
.quiz-question-text {{
    font-size: 1.05rem; font-weight: 600; color: {TEXT} !important;
    line-height: 1.55; margin-bottom: 0.5rem;
}}
.quiz-result-correct {{
    background: #F0FDF4; border: 2px solid #4ADE80; border-radius: 8px;
    padding: 1rem 1.2rem; margin-top: 0.8rem; color: #14532D !important;
}}
.quiz-result-wrong {{
    background: #FFF1F2; border: 2px solid #F87171; border-radius: 8px;
    padding: 1rem 1.2rem; margin-top: 0.8rem; color: #7F1D1D !important;
}}
.quiz-explanation {{
    background: #EFF6FF; border-left: 4px solid {BLUE}; border-radius: 0 8px 8px 0;
    padding: 1rem 1.2rem; margin-top: 0.8rem; font-size: 0.92rem;
    line-height: 1.65; color: {TEXT} !important;
}}

/* Mode selector cards */
.mode-card-active {{
    background: #EEF2FF; border: 3px solid {BLUE}; border-radius: 12px;
    padding: 1.2rem 1.4rem; margin-bottom: 0.6rem;
    box-shadow: 0 4px 12px rgba(0,48,135,0.15); color: {TEXT};
}}
.mode-card-inactive {{
    background: {CARD}; border: 2px solid {BORDER}; border-radius: 12px;
    padding: 1.2rem 1.4rem; margin-bottom: 0.6rem;
    box-shadow: 0 2px 6px rgba(0,48,135,0.06); color: {TEXT};
}}

/* Pills */
.pill-ok {{
    display: inline-block; background: #DCFCE7; color: #166534;
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid #4ADE80;
}}
.pill-warn {{
    display: inline-block; background: #FEF3C7; color: #92400E;
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid #FCD34D;
}}
.pill-err {{
    display: inline-block; background: #FEE2E2; color: #991B1B;
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid #F87171;
}}
.pill-blue {{
    display: inline-block; background: #EEF2FF; color: {BLUE};
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid {BORDER};
}}

/* Misc */
.streamlit-expanderHeader {{
    background-color: #EEF2FF !important; color: {BLUE} !important;
    font-weight: 600 !important; border-radius: 8px !important;
}}
.rb-footer {{
    text-align: center; color: {MUTED}; font-size: 0.78rem;
    border-top: 1px solid {BORDER}; padding-top: 1rem; margin-top: 2.5rem;
}}
.ref-log {{
    background: #EEF2FF; border: 1px solid {BORDER};
    border-left: 4px solid {BLUE}; border-radius: 0 8px 8px 0;
    padding: 0.9rem 1.1rem; font-size: 0.88rem; color: {TEXT};
}}
#MainMenu {{ display: none !important; }}
[data-testid="stMainMenu"] {{ display: none !important; }}
footer {{ visibility: hidden; }}
/* Transparent header — blends with cream page background, no black bar */
[data-testid="stHeader"] {{ background: transparent !important; }}
/* Hide the Deploy button and its container */
[data-testid="stAppDeployButton"], .stAppDeployButton {{ display: none !important; }}
/* Keep only the Share button — hide every sibling after the first one in the
   toolbar actions area (GitHub source, fork/star, and kebab icons).
   Share is always the first child so :nth-child(n+2) targets the rest. */
[data-testid="stToolbarActions"] > *:nth-child(n+2) {{ display: none !important; }}
[data-testid="stSlider"] .st-by {{ background: {BLUE} !important; }}
[data-baseweb="select"] {{ background-color: {CARD} !important; }}
.stAlert {{ border-radius: 8px !important; font-size: 0.88rem !important; }}

/* Inputs */
.stTextArea textarea, .stTextInput input {{
    background-color: {CARD} !important; color: {TEXT} !important;
    border: 1.5px solid {BORDER} !important; border-radius: 8px !important;
    font-size: 0.92rem !important;
}}
.stTextArea textarea:focus, .stTextInput input:focus {{
    border-color: {BLUE} !important; box-shadow: 0 0 0 3px rgba(0,48,135,0.12) !important;
}}
[data-testid="stFileUploader"] {{
    border: 2px dashed {BLUE_L} !important; border-radius: 10px !important;
    background-color: #EEF2FF !important; padding: 0.5rem;
}}
.stChatMessage {{
    background: {CARD} !important; border: 1px solid {BORDER} !important;
    border-radius: 10px !important; margin-bottom: 0.5rem;
    box-shadow: 0 1px 4px rgba(0,48,135,0.06);
}}

/* Accuracy bar */
.accuracy-bar-wrap {{
    background: #E2E8F0; border-radius: 20px; height: 10px;
    margin: 6px 0 2px 0; overflow: hidden;
}}
.accuracy-bar-fill {{
    height: 10px; border-radius: 20px;
    background: linear-gradient(90deg, {BLUE} 0%, {BLUE_L} 100%);
    transition: width 0.4s ease;
}}

/* ── Multiselect chips (selected options) ────────────────────────────────────
   Streamlit renders each selected multiselect option inside
   [data-testid="stMultiSelectTagsContainer"] as an element carrying a bare
   `data-tag` attribute, with the remove control labelled aria-label="Remove X".
   It is NOT a BaseWeb [data-baseweb="tag"] element in this version — selectors
   written against that name match nothing and silently do nothing.

   The chip background comes from primaryColor (navy), while the earlier blanket
   rule `[data-baseweb="select"] span, [data-baseweb="select"] div {{ color: #1F2937 }}`
   also matches the text inside the chip — producing near-black text on navy.

   These rules use a descendant pair, e.g.
   [data-testid="stMultiSelectTagsContainer"] [data-tag] *  →  specificity (0,2,0),
   which outranks the blanket rule's (0,1,1), and are declared last so document
   order backs them up too. -webkit-text-fill-color is set alongside `color`
   because in WebKit (Safari, Chrome on iOS) it overrides `color` regardless of
   specificity, so setting `color` alone can silently fail on a phone. */

/* The chip itself — navy background AND white text set together here.
   CRITICAL: the transparent-background rule below must NOT list the bare
   [data-tag] selector. Both rules carry specificity (0,2,0), so if the chip
   element appears in both, the later one wins and paints the chip transparent —
   revealing the near-white container behind it, leaving white text on white. */
[data-testid="stMultiSelectTagsContainer"] [data-tag],
[data-testid="stMultiSelect"] [data-tag],
[data-baseweb="tag"] {{
    background-color: {BLUE} !important;
    border: 1px solid {BLUE} !important;
    border-radius: 6px !important;
    color: #FFFFFF !important;
    -webkit-text-fill-color: #FFFFFF !important;
}}

/* Chip DESCENDANTS only ( * ) — white text, and transparent backgrounds so the
   inner label/button don't paint over the chip's navy. */
[data-testid="stMultiSelectTagsContainer"] [data-tag] *,
[data-testid="stMultiSelect"] [data-tag] *,
[data-baseweb="tag"] * {{
    color: #FFFFFF !important;
    -webkit-text-fill-color: #FFFFFF !important;
    background-color: transparent !important;
    background: transparent !important;
    font-weight: 600 !important;
}}

/* The "x" remove control on each chip */
[data-testid="stMultiSelectTagsContainer"] [data-tag] [aria-label^="Remove"],
[data-testid="stMultiSelectTagsContainer"] [data-tag] svg,
[data-testid="stMultiSelectTagsContainer"] [data-tag] path,
[data-testid="stMultiSelect"] [data-tag] [aria-label^="Remove"],
[data-testid="stMultiSelect"] [data-tag] svg,
[data-testid="stMultiSelect"] [data-tag] path {{
    color: #FFFFFF !important;
    fill: #FFFFFF !important;
    stroke: #FFFFFF !important;
    background: transparent !important;
    opacity: 0.85;
}}
[data-testid="stMultiSelectTagsContainer"] [data-tag] [aria-label^="Remove"]:hover,
[data-testid="stMultiSelect"] [data-tag] [aria-label^="Remove"]:hover {{
    opacity: 1;
    background-color: rgba(255,255,255,0.25) !important;
    border-radius: 3px;
}}
</style>
""", unsafe_allow_html=True)



# =============================================================================
# CONFIG + ACCESS GATE
#
# Two supported ways to supply configuration:
#
#   Render / Docker / any normal host  → environment variables
#       ANTHROPIC_API_KEY = sk-ant-...
#       APP_PASSWORD      = your-crew-password
#
#   Local dev / Streamlit Cloud        → .streamlit/secrets.toml
#       ANTHROPIC_API_KEY = "sk-ant-..."
#       APP_PASSWORD      = "your-crew-password"
#
# If APP_PASSWORD is NOT set, the gate is skipped entirely (convenient for local
# development). NEVER deploy publicly without setting APP_PASSWORD.
# =============================================================================

def get_secret(name: str, default=None):
    """
    Read a configuration value from, in order:

      1. OS environment variables  — Render, Docker, Fly, any normal host
      2. Streamlit secrets.toml    — local dev, Streamlit Community Cloud
      3. Nested [anthropic] table  — legacy secrets.toml layout

    IMPORTANT: st.secrets does NOT read arbitrary environment variables. If no
    secrets.toml file exists anywhere on disk, merely *touching* st.secrets
    raises StreamlitSecretNotFoundError — which is not a KeyError, so a narrow
    `except KeyError` will not catch it and the app crashes on boot. That is
    why every st.secrets access below is wrapped in a broad `except Exception`.
    """
    # 1. Environment variable
    val = os.environ.get(name)
    if val:
        return val

    # 2. Flat key in secrets.toml
    try:
        val = st.secrets[name]
        if val:
            return val
    except Exception:
        pass

    # 3. Legacy nested table, e.g. [anthropic] api_key = "..."
    if name == "ANTHROPIC_API_KEY":
        try:
            val = st.secrets["anthropic"]["api_key"]
            if val:
                return val
        except Exception:
            pass

    return default


def running_on_render() -> bool:
    """Render sets RENDER=true in every service environment."""
    return os.environ.get("RENDER", "").lower() in ("true", "1", "yes")


def check_password() -> bool:
    """
    Render a password prompt and return True only once the correct password
    has been entered. Uses hmac.compare_digest for constant-time comparison
    so the check isn't vulnerable to timing attacks.
    """
    expected = get_secret("APP_PASSWORD")

    if not expected:
        # No password configured.
        if running_on_render():
            # FAIL CLOSED. A public host with no password would expose the
            # owner's API key to the open internet — never allow that silently.
            st.error(
                "🔒 **APP_PASSWORD is not set.**\n\n"
                "This app is running on a public host with no access control, "
                "so it has been locked as a precaution.\n\n"
                "Fix: Render Dashboard → your service → **Environment** → "
                "add `APP_PASSWORD`, then **Save Changes**."
            )
            st.stop()
        # Local development — no password needed, allow through.
        return True

    # Already authenticated this session
    if st.session_state.get("_auth_ok", False):
        return True

    # ── Login screen ─────────────────────────────────────────────────────────
    st.markdown("""
    <div style="text-align:center;padding:3rem 2rem 1rem 2rem;">
        <div style="color:#003087;font-size:3rem;font-weight:900;
                    letter-spacing:-1.5px;">🏀 RefBuddy</div>
        <div style="color:#4B5563;font-size:1.1rem;margin-bottom:0.5rem;">
            Built by a ref, for refs</div>
        <div style="color:#6B7280;font-size:0.9rem;">
            Private tool for MSHSL officials — enter your crew password to continue.</div>
    </div>
    """, unsafe_allow_html=True)

    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        pw = st.text_input("Password", type="password",
                           key="_pw_input", label_visibility="collapsed",
                           placeholder="Crew password")
        if st.button("Enter", use_container_width=True, key="_pw_submit"):
            if hmac.compare_digest(str(pw), str(expected)):
                st.session_state["_auth_ok"] = True
                st.rerun()
            else:
                st.error("❌ Incorrect password.")
        st.caption("Need access? Contact your assignor or crew chief.")

    return False


if not check_password():
    st.stop()


# =============================================================================
# SESSION USAGE CAP
# Second layer of cost protection: caps how many analysis frames one browser
# session can send to the API. Film analysis is by far the most expensive
# operation (~1,500 tokens per frame), so this is where the money goes.
# =============================================================================

MAX_FRAMES_PER_SESSION = 400   # ~600K tokens ≈ a few dollars, generous for one sitting


def frames_budget_left() -> int:
    """How many analysis frames this session has left."""
    used = st.session_state.get("frames_used", 0)
    return max(0, MAX_FRAMES_PER_SESSION - used)


def spend_frames(n: int) -> bool:
    """
    Charge n frames against the session budget.
    Returns False (and shows an error) if the budget is exhausted.
    """
    if frames_budget_left() < n:
        st.error(
            f"⚠️ **Session analysis limit reached** "
            f"({MAX_FRAMES_PER_SESSION} frames). "
            "Refresh the page to start a new session, or reduce your frame range. "
            "This limit keeps API costs predictable."
        )
        return False
    st.session_state["frames_used"] = st.session_state.get("frames_used", 0) + n
    return True


# =============================================================================
# SESSION STATE
# =============================================================================

def _s(k, v):
    if k not in st.session_state:
        st.session_state[k] = v

_s("messages", [])
_s("uploaded_files_content", [])

# Session usage budget
_s("frames_used", 0)

# Film & Grade (merged — photos + video share one path)
_s("fg_frames", [])
_s("fg_labels", [])
_s("fg_source", "")
_s("fg_is_video", False)
_s("fg_fps", 1.0)
_s("fg_result", "")
_s("rg_saved_logs", [])

# Ref Hub
_s("ah_sub", "pregame")
_s("ah_eval_result", "")
_s("ah_eval_scope", "")
_s("ah_pregame_result", "")
_s("ah_pregame_logs", [])

# Quiz
_s("quiz_mode", None)
_s("quiz_topic", "Mixed")
_s("quiz_current_q", None)
_s("quiz_answered", False)
_s("quiz_user_answer", None)
_s("quiz_total", 0)
_s("quiz_correct", 0)
_s("quiz_session_topics", [])
_s("tenq_questions", [])
_s("tenq_index", 0)
_s("tenq_answers", [])
_s("tenq_finished", False)
_s("tenq_answered_this", False)
_s("tenq_user_answer", None)
_s("quiz_log", [])


# =============================================================================
# HELPERS — Core
# =============================================================================

def b64(data: bytes) -> str:
    return base64.standard_b64encode(data).decode("utf-8")

MODEL = "claude-sonnet-4-6"

# =============================================================================
# ANTHROPIC SDK COMPATIBILITY SHIM
#
# INCIDENT: anthropic 1.0.0 (a major version bump) REMOVED `temperature`,
# `top_p`, and `top_k` from the typed signatures of both messages.create() and
# messages.stream(). Because requirements.txt originally pinned nothing, a
# routine Render redeploy resolved to the new major version and every API call
# in the app began raising:
#
#     TypeError: Messages.stream() got an unexpected keyword argument 'temperature'
#
# The sampling controls did not simply move — the new OutputConfigParam carries
# only `effort` and `format`, and there is no **kwargs passthrough to absorb the
# argument.
#
# The real fix is the upper bounds now in requirements.txt. This shim is
# insurance: it introspects the INSTALLED SDK at runtime and passes temperature
# only if that version accepts it. If a future SDK drops it again, the app
# degrades to default sampling instead of going down mid-season.
#
# NOTE ON temperature=0: it matters here. RefBuddy is a rules reference — the
# same question about a shot clock reset should return the same answer every
# time. If you ever deliberately move past anthropic 1.0.0, this shim keeps the
# app running but you will be on the API's default sampling. Revisit the
# determinism approach at that point rather than letting it change silently.
# =============================================================================

@functools.lru_cache(maxsize=2)
def _sdk_accepts(param: str) -> bool:
    """Does the installed anthropic SDK accept this sampling parameter?"""
    try:
        import inspect
        client = anthropic.Anthropic(api_key="sk-ant-introspection-only")
        for fn in (client.messages.create, client.messages.stream):
            params = inspect.signature(fn).parameters
            if param not in params and not any(
                p.kind == p.VAR_KEYWORD for p in params.values()
            ):
                return False
        return True
    except Exception:
        # If introspection fails for any reason, omit the parameter rather than
        # risk a TypeError that takes down every request.
        return False


def temp_kwargs() -> dict:
    """
    Returns {"temperature": 0} on SDKs that support it, {} on those that don't.
    Spread into every API call with **temp_kwargs().
    """
    return {"temperature": 0} if _sdk_accepts("temperature") else {}


def make_client():
    """
    Create an Anthropic client using get_secret(), which checks environment
    variables first (Render) and secrets.toml second (local dev).
    """
    key = get_secret("ANTHROPIC_API_KEY")
    if not key:
        st.error(
            "❌ **ANTHROPIC_API_KEY not found.**\n\n"
            "**On Render:** Dashboard → your service → **Environment** → add "
            "`ANTHROPIC_API_KEY`, then **Save Changes**.\n\n"
            "**Locally:** add it to `.streamlit/secrets.toml`:\n```\n"
            'ANTHROPIC_API_KEY = "sk-ant-..."\n```'
        )
        st.stop()
    return anthropic.Anthropic(api_key=key)

def api_key_ok() -> bool:
    return True

def handle_api_error(e: Exception) -> str:
    if isinstance(e, anthropic.AuthenticationError):
        return "❌ Authentication failed. Check ANTHROPIC_API_KEY in Streamlit secrets."
    if isinstance(e, anthropic.RateLimitError):
        return "⚠️ Rate limit reached. Wait a moment and try again."
    if isinstance(e, anthropic.APIConnectionError):
        return "❌ Connection error. Check your internet connection."
    if isinstance(e, anthropic.BadRequestError):
        return (f"❌ Request too large or malformed: {e}\n\n"
                "Try reducing the frame range or using 0.5 fps extraction.")
    return f"❌ Unexpected error: {e}"

def prepare_file_content(uf):
    data = uf.read()
    name = uf.name.lower()
    if name.endswith(".pdf"):
        return {"type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": b64(data)},
                "title": uf.name}
    if name.endswith((".jpg", ".jpeg")):
        return {"type": "image",
                "source": {"type": "base64", "media_type": "image/jpeg", "data": b64(data)}}
    if name.endswith(".png"):
        return {"type": "image",
                "source": {"type": "base64", "media_type": "image/png", "data": b64(data)}}
    if name.endswith(".txt"):
        return {"type": "text",
                "text": f"[File: {uf.name}]\n\n{data.decode('utf-8', errors='replace')}"}
    return None

def stream_chat(client, messages, files, system=None):
    sys_p = system or SYSTEM_PROMPT
    api_msgs = []
    for i, m in enumerate(messages):
        if m["role"] == "user" and i == len(messages) - 1 and files:
            blocks = list(files) + [{"type": "text", "text": m["content"]}]
            api_msgs.append({"role": "user", "content": blocks})
        else:
            api_msgs.append({"role": m["role"], "content": m["content"]})
    with client.messages.stream(
        model=MODEL, max_tokens=8192,   # raised from 4096: long rules
        # breakdowns and multi-part answers were at risk of stopping mid-sentence
        system=system_blocks(sys_p), messages=api_msgs, **temp_kwargs(),
    ) as s:
        yield from s.text_stream

def call_api_sync(prompt: str, system: str, max_tokens: int = 3000) -> str:
    client = make_client()
    resp = client.messages.create(
        model=MODEL, max_tokens=max_tokens,
        system=system_blocks(system), messages=[{"role": "user", "content": prompt}], **temp_kwargs(),
    )
    return resp.content[0].text

def _human_ts(iso: str) -> str:
    """Turn a stored ISO timestamp into 'Aug 24, 2026 at 07:24 PM'."""
    if not iso:
        return ""
    try:
        return datetime.datetime.fromisoformat(iso).strftime("%b %d, %Y at %I:%M %p")
    except Exception:
        return iso[:19]


def chat_log_markdown() -> str:
    """
    Single source of truth for the TXT, PDF, and Word chat-log exports, so all
    three contain identical content.

    Questions and answers are PAIRED under numbered headings rather than dumped
    as a flat role/content list — a flat list is what made exports read as
    truncated blobs even when the full answer was present.
    """
    msgs = st.session_state.messages
    now = datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")

    out = ["RefBuddy Chat Log",
           "=================",
           f"Exported: {now}",
           f"Messages: {len(msgs)}",
           "_" * 70,
           ""]

    qnum = 0
    i = 0
    while i < len(msgs):
        m = msgs[i]
        if m["role"] == "user":
            qnum += 1
            out.append(f"Question {qnum}")
            out.append("-" * len(f"Question {qnum}"))
            out.append(_human_ts(m.get("timestamp", "")))
            out.append("")
            out.append(m["content"].strip())
            out.append("")
            # Pair the assistant reply that follows, if there is one
            if i + 1 < len(msgs) and msgs[i + 1]["role"] == "assistant":
                out.append("RefBuddy Response")
                out.append("")
                out.append(msgs[i + 1]["content"].strip())
                out.append("")
                i += 1
            else:
                out.append("RefBuddy Response")
                out.append("")
                out.append("_(No response recorded for this question.)_")
                out.append("")
            out.append("_" * 70)
            out.append("")
        else:
            # An assistant message with no preceding question (shouldn't happen,
            # but never silently drop content from an export)
            out.append("RefBuddy Response")
            out.append("")
            out.append(m["content"].strip())
            out.append("")
            out.append("_" * 70)
            out.append("")
        i += 1

    out.append("")
    out.append("Always confirm rulings with your MSHSL assignor. "
               "Not official NFHS/MSHSL interpretation.")
    return "\n".join(out)


def chat_log_json() -> str:
    """Structured export. Roles read as You / RefBuddy; timestamps humanized."""
    return json.dumps({
        "exported_at": datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p"),
        "message_count": len(st.session_state.messages),
        "messages": [
            {
                "role": "You" if m["role"] == "user" else "RefBuddy",
                "timestamp": _human_ts(m.get("timestamp", "")),
                "content": m["content"],
            }
            for m in st.session_state.messages
        ],
        "disclaimer": ("Always confirm rulings with your MSHSL assignor. "
                       "Not official NFHS/MSHSL interpretation."),
    }, indent=2, ensure_ascii=False)


# =============================================================================
# HELPERS — Frame extraction
# =============================================================================

def _asset_data_uri(filename: str) -> str | None:
    """
    Load a bundled image (e.g. Claude.png) from the app directory and return it
    as a base64 data URI for inline <img> use. Returns None if the file is
    missing so the caller can fall back to text — a missing logo should never
    break the page. Path is resolved relative to app.py, so it behaves the same
    locally and on Render.
    """
    try:
        here = os.path.dirname(os.path.abspath(__file__))
        path = os.path.join(here, filename)
        with open(path, "rb") as f:
            return "data:image/png;base64," + base64.standard_b64encode(f.read()).decode()
    except Exception:
        return None


def image_to_frame_b64(uploaded_image) -> str | None:
    """
    Convert an uploaded still image (JPG/PNG) into the same base64-JPEG format
    extract_frames() produces, so photos and video frames flow through one
    identical downstream path.

    Resized to max 1280px wide to control token cost, same as video frames.
    """
    try:
        data = uploaded_image.read()
        arr = np.frombuffer(data, np.uint8)
        img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if img is None:
            return None
        h, w = img.shape[:2]
        if w > 1280:
            img = cv2.resize(img, (1280, int(h * 1280 / w)), interpolation=cv2.INTER_AREA)
        ok, buf = cv2.imencode(".jpg", img, [cv2.IMWRITE_JPEG_QUALITY, 88])
        return base64.standard_b64encode(buf).decode("utf-8") if ok else None
    except Exception:
        return None


def extract_frames(video_path: str, fps: float = 1.0) -> list:
    """Extract frames from video at specified fps. Returns list of base64 JPEG strings."""
    if not OPENCV_AVAILABLE:
        raise RuntimeError("opencv-python-headless not available.")
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise RuntimeError(f"Cannot open video: {video_path}")
    native_fps = cap.get(cv2.CAP_PROP_FPS) or 30.0
    interval = max(1, int(round(native_fps / fps)))
    frames, idx = [], 0
    while True:
        ret, frame = cap.read()
        if not ret:
            break
        if idx % interval == 0:
            h, w = frame.shape[:2]
            if w > 1280:
                frame = cv2.resize(frame, (1280, int(h * 1280 / w)), interpolation=cv2.INTER_AREA)
            ok, buf = cv2.imencode(".jpg", frame, [cv2.IMWRITE_JPEG_QUALITY, 85])
            if ok:
                frames.append(base64.standard_b64encode(buf).decode("utf-8"))
        idx += 1
    cap.release()
    return frames

def extract_video_uploaded(uploaded_video, fps: float = 1.0) -> tuple:
    suffix = ".mp4" if uploaded_video.name.lower().endswith(".mp4") else ".mov"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_video.read())
        tmp_path = tmp.name
    frames = extract_frames(tmp_path, fps=fps)
    os.unlink(tmp_path)
    return frames, uploaded_video.name

def build_vision_content(frames_b64, start_idx, end_idx, user_question,
                          video_name, fps_used, preamble_extra="") -> list:
    selected = frames_b64[start_idx: end_idx + 1]
    spf = 1.0 / fps_used
    content = [{"type": "text", "text": (
        f"Game film: {video_name}\n"
        f"Frames: {len(selected)} ({start_idx+1}–{end_idx+1} of {len(frames_b64)}) "
        f"at {fps_used} fps ({spf:.1f}s/frame).\n"
        f"Frame numbering is 1-based. Use 'Frame N' format.\n{preamble_extra}\n"
    )}]
    for i, fb in enumerate(selected):
        fn = start_idx + i + 1
        content.append({"type": "text", "text": f"--- Frame {fn} (~{(fn-1)/fps_used:.1f}s) ---"})
        content.append({"type": "image", "source": {"type": "base64",
                         "media_type": "image/jpeg", "data": fb}})
    content.append({"type": "text", "text": f"\nQuestion:\n{user_question}"})
    return content

def stream_vision(client, content_blocks, system):
    with client.messages.stream(
        model=MODEL, max_tokens=8192, system=system_blocks(system),   # raised from 4096
        messages=[{"role": "user", "content": content_blocks}], **temp_kwargs(),
    ) as s:
        yield from s.text_stream


# =============================================================================
# HELPERS — Quiz engine
# =============================================================================

def _strip_json_fences(raw: str) -> str:
    raw = raw.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return raw.strip()

def generate_single_question(topic: str, used_topics: list = None) -> dict | None:
    if not api_key_ok():
        return None
    avoid_str = ""
    if used_topics and len(used_topics) > 0:
        recent = used_topics[-5:]
        avoid_str = (f"IMPORTANT: Do NOT generate a question about any of these topics "
                     f"that were just asked: {', '.join(recent)}. "
                     "Pick a completely different rule, mechanic, or scenario.\n")
    import random
    q_type_hint = "true_false" if random.random() < 0.5 else "multiple_choice"
    topic_str = "" if topic == "Mixed" else f"Topic focus: {topic}. "
    prompt = (
        f"{avoid_str}"
        f"{topic_str}"
        f"Generate one {q_type_hint} question for a Minnesota high school basketball referee. "
        f"It must be challenging, specific, and reference exact rule numbers or MSHSL modifications. "
        f"For multiple_choice: EXACTLY 4 options (A, B, C, D). "
        f"For true_false: EXACTLY 2 options (A=True, B=False). "
        f"Respond with ONLY valid JSON — no fences, no preamble."
    )
    try:
        client = make_client()
        resp = client.messages.create(
            model=MODEL, max_tokens=900,
            system=system_blocks(QUIZ_SYSTEM_PROMPT),
            messages=[{"role": "user", "content": prompt}],
            **temp_kwargs(),
        )
        raw = _strip_json_fences(resp.content[0].text)
        q = json.loads(raw)
        if q.get("type") == "multiple_choice" and len(q.get("options", {})) != 4:
            return None
        if q.get("type") == "true_false" and len(q.get("options", {})) != 2:
            return None
        return q
    except Exception as e:
        st.error(f"❌ Failed to generate question: {e}")
        return None

def generate_ten_questions(topic: str) -> list | None:
    if not api_key_ok():
        return None
    topic_str = "" if topic == "Mixed" else f"Topic focus: {topic}. "
    prompt = (
        f"{topic_str}Generate exactly 10 questions for a Minnesota high school basketball referee. "
        "Mix: 5 multiple_choice (EXACTLY 4 options A/B/C/D each) + 5 true_false. "
        "Cover these areas: 2025-26 rule changes (no offensive goaltending, backboard slap), "
        "MSHSL mercy rule, MSHSL bonus system (1&1 half-based), shot clock resets (35/20), "
        "restricted area arc, flopping mechanics, Lead/Trail/Center positioning, "
        "throw-in spots (3-pt line rule), personal game note scenarios, closely guarded dribbling exception. "
        "Respond with ONLY a valid JSON array of 10 objects — no fences, no preamble."
    )
    try:
        client = make_client()
        resp = client.messages.create(
            model=MODEL, max_tokens=6000,
            system=system_blocks(QUIZ_SYSTEM_PROMPT),
            messages=[{"role": "user", "content": prompt}],
            **temp_kwargs(),
        )
        raw = _strip_json_fences(resp.content[0].text)
        questions = json.loads(raw)
        if isinstance(questions, list) and len(questions) == 10:
            return questions
        if isinstance(questions, list) and len(questions) >= 5:
            return questions[:10]
        st.error("❌ Unexpected question count. Try again.")
        return None
    except Exception as e:
        st.error(f"❌ Failed to generate quiz: {e}")
        return None

def render_question_card(q: dict, question_num: str = ""):
    q_text = q.get("question", "")
    q_type = q.get("type", "multiple_choice")
    badge_color = "#EEF2FF"
    badge_label = "True/False" if q_type == "true_false" else "Multiple Choice"
    st.markdown(f"""
    <div class="quiz-question-card">
        <div class="quiz-question-text">{question_num} {q_text}
        <span style="background:{badge_color};color:{BLUE};font-size:0.72rem;
        font-weight:700;border-radius:20px;padding:2px 8px;margin-left:8px;">
        {badge_label}</span></div>
    </div>
    """, unsafe_allow_html=True)

def render_feedback(q: dict, user_answer: str) -> bool:
    correct = q.get("correct", "")
    options = q.get("options", {})
    correct_text = options.get(correct, correct)
    user_text = options.get(user_answer, user_answer)
    is_correct = user_answer == correct
    if is_correct:
        st.markdown(f"""<div class="quiz-result-correct">
        <strong>✅ Correct!</strong> &nbsp; {user_answer}: {user_text}
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown(f"""<div class="quiz-result-wrong">
        <strong>❌ Incorrect.</strong> You chose: {user_answer}: {user_text}<br>
        <strong>✔ Correct: {correct}: {correct_text}</strong>
        </div>""", unsafe_allow_html=True)
    explanation = q.get("explanation", "")
    rule_cite = q.get("rule_citation", "")
    personal = q.get("personal_note", "")
    pnote = f'<br><strong>📋 From RefBuddy Knowledge Base:</strong> {personal}' if personal else ""
    st.markdown(f"""<div class="quiz-explanation">
    <strong>📖 Explanation</strong><br>{explanation}<br><br>
    <strong>📌 Citation:</strong> {rule_cite}{pnote}
    </div>""", unsafe_allow_html=True)
    return is_correct

def accuracy_display(correct: int, total: int):
    pct = int(round(correct / total * 100)) if total > 0 else 0
    color = "#15803D" if pct >= 80 else ("#92400E" if pct >= 60 else "#991B1B")
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:12px;background:{CARD};
                border:1px solid {BORDER};border-radius:8px;padding:0.7rem 1rem;
                margin-bottom:0.8rem;">
        <div style="font-weight:800;font-size:1.4rem;color:{color};min-width:52px;">{pct}%</div>
        <div style="flex:1;">
            <div class="accuracy-bar-wrap">
                <div class="accuracy-bar-fill" style="width:{pct}%;background:{color};"></div>
            </div>
            <div style="font-size:0.8rem;color:{MUTED};margin-top:3px;">
                {correct} correct of {total} answered</div>
        </div>
    </div>""", unsafe_allow_html=True)


# =============================================================================
# HELPERS — Export (PDF + DOCX)
# =============================================================================

def sanitize_for_pdf(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "\u2014": "-", "\u2013": "-", "\u2012": "-", "\u2015": "-",
        "\u2018": "'", "\u2019": "'", "\u201a": ",", "\u201b": "'",
        "\u201c": '"', "\u201d": '"', "\u201e": '"', "\u201f": '"',
        "\u2026": "...", "\u2022": "*", "\u2023": ">", "\u25e6": "o", "\u2043": "-",
        "\u00a0": " ", "\u200b": "", "\u200c": "", "\u200d": "", "\u2060": "", "\ufeff": "",
        "\u00d7": "x", "\u00f7": "/", "\u2212": "-", "\u00b0": "°",
        "\u2192": "->", "\u2190": "<-", "\u2194": "<->",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text.strip()

def markdown_to_pdf_bytes(md_text: str, title: str = "RefBuddy Report") -> bytes | None:
    try:
        from fpdf import FPDF
        L_MARGIN = 20; R_MARGIN = 20; TOP_MARGIN = 18; BOT_MARGIN = 15
        pdf = FPDF()
        pdf.set_margins(L_MARGIN, TOP_MARGIN, R_MARGIN)
        pdf.set_auto_page_break(auto=True, margin=BOT_MARGIN)
        pdf.add_page()
        eff_w = pdf.w - L_MARGIN - R_MARGIN
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(0, 48, 135)
        pdf.multi_cell(eff_w, 7, sanitize_for_pdf(title), align="L")
        pdf.ln(1)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(107, 114, 128)
        date_str = sanitize_for_pdf(
            f"Generated: {datetime.datetime.now().strftime('%B %d, %Y %H:%M')}"
        )
        pdf.multi_cell(eff_w, 4, date_str, align="L")
        pdf.ln(2)
        pdf.set_draw_color(180, 200, 220)
        pdf.set_line_width(0.3)
        pdf.line(L_MARGIN, pdf.get_y(), pdf.w - R_MARGIN, pdf.get_y())
        pdf.ln(3)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(31, 41, 55)
        for raw_line in md_text.split("\n"):
            s = raw_line.strip()
            if not s:
                pdf.ln(2); continue
            if s == "---":
                pdf.ln(1)
                pdf.set_draw_color(200, 210, 220)
                pdf.line(L_MARGIN, pdf.get_y(), pdf.w - R_MARGIN, pdf.get_y())
                pdf.ln(2); continue
            if s.startswith("# ") or s.startswith("## "):
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(0, 48, 135)
                pdf.multi_cell(eff_w, 6, sanitize_for_pdf(s.lstrip("#").strip()), align="L")
                pdf.set_font("Helvetica", "", 9); pdf.set_text_color(31, 41, 55); continue
            if s.startswith("### "):
                pdf.ln(1)
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(0, 48, 135)
                pdf.multi_cell(eff_w, 5, sanitize_for_pdf(s.lstrip("#").strip()), align="L")
                pdf.set_font("Helvetica", "", 9); pdf.set_text_color(31, 41, 55); continue
            if s.startswith("#### "):
                pdf.ln(1)
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_text_color(31, 41, 55)
                pdf.multi_cell(eff_w, 5, sanitize_for_pdf(s.lstrip("#").strip()), align="L")
                pdf.set_font("Helvetica", "", 9); continue
            if s.startswith(("- ", "* ", "+ ")):
                content = sanitize_for_pdf(s[2:].replace("**", ""))
                pdf.set_x(L_MARGIN); pdf.cell(6, 4.5, "-")
                pdf.set_x(L_MARGIN + 6); pdf.multi_cell(eff_w - 6, 4.5, content, align="L"); continue
            import re as _re
            num_match = _re.match(r"^(\d+)\.\s+(.*)", s)
            if num_match:
                num = num_match.group(1) + "."; content = sanitize_for_pdf(num_match.group(2).replace("**", ""))
                pdf.set_x(L_MARGIN); pdf.cell(8, 4.5, num)
                pdf.set_x(L_MARGIN + 8); pdf.multi_cell(eff_w - 8, 4.5, content, align="L"); continue
            if s.startswith("| ") or (s.startswith("|") and "|" in s[1:]):
                stripped = s.replace("|", "").replace("-", "").replace(" ", "")
                if not stripped: continue
                row = sanitize_for_pdf(s.strip("|").replace("|", "  ").replace("**", ""))
                pdf.set_font("Courier", "", 8); pdf.multi_cell(eff_w, 4, row, align="L")
                pdf.set_font("Helvetica", "", 9); continue
            clean = sanitize_for_pdf(s.replace("**", "").replace("*", ""))
            pdf.multi_cell(eff_w, 4.5, clean, align="L")
        return bytes(pdf.output())
    except ImportError:
        return None

def markdown_to_docx_bytes(md_text: str, title: str = "RefBuddy Report") -> bytes | None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1); section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2); section.right_margin = Inches(1.2)
        t_para = doc.add_paragraph()
        t_run = t_para.add_run(title)
        t_run.bold = True; t_run.font.size = Pt(18); t_run.font.color.rgb = RGBColor(0, 48, 135)
        t_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        d_para = doc.add_paragraph()
        d_run = d_para.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y %H:%M')}")
        d_run.font.size = Pt(9); d_run.font.color.rgb = RGBColor(107, 114, 128)
        doc.add_paragraph()
        for line in md_text.split("\n"):
            s = line.strip()
            if s.startswith("## ") or s.startswith("# "):
                h = doc.add_heading(s.lstrip("#").strip(), level=2)
                for run in h.runs: run.font.color.rgb = RGBColor(0, 48, 135)
            elif s.startswith("### "):
                h = doc.add_heading(s.lstrip("#").strip(), level=3)
                for run in h.runs: run.font.color.rgb = RGBColor(0, 48, 135)
            elif s.startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(s[2:].replace("**", "")).font.size = Pt(10)
            elif s == "---":
                doc.add_paragraph(); doc.add_paragraph().add_run("─" * 60).font.size = Pt(8); doc.add_paragraph()
            elif s == "":
                doc.add_paragraph()
            else:
                p = doc.add_paragraph(); p.add_run(s.replace("**", "")).font.size = Pt(10)
        buf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx"); buf.close()
        doc.save(buf.name)
        with open(buf.name, "rb") as f: data = f.read()
        os.unlink(buf.name)
        return data
    except ImportError:
        return None


# =============================================================================
# SIDEBAR
# =============================================================================

with st.sidebar:
    st.markdown(
        '<div style="background:#F8FAFC;border:2px solid #1F2937;border-radius:8px;'
        'padding:0.7rem 1rem;margin-bottom:0.8rem;">'
        '<span style="color:#1F2937;font-weight:800;font-size:1.1rem;">🏀 RefBuddy</span><br>'
        '<span style="color:#4B5563;font-size:0.72rem;">Built by a ref, for refs</span>'
        '</div>',
        unsafe_allow_html=True,
    )

    # "Powered by Claude" lockup — base64-embedded so it works identically
    # locally and on Render. Falls back to plain text if the file is missing.
    _claude_uri = _asset_data_uri("Claude.png")
    if _claude_uri:
        st.markdown(
            '<div style="display:flex;align-items:center;gap:12px;'
            'flex-wrap:wrap;margin:0.7rem 0 0.4rem 0;">'
            '<span style="color:#1F2937;font-weight:700;font-size:1.05rem;'
            'line-height:1;">Powered by</span>'
            f'<img src="{_claude_uri}" alt="Claude" '
            'style="height:85px;width:auto;display:block;">'
            '</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<div style="color:#1F2937;font-weight:700;font-size:0.9rem;'
            'margin:0.5rem 0 0.2rem 0;">Powered by Claude</div>',
            unsafe_allow_html=True,
        )

    # Session analysis budget indicator
    _left = frames_budget_left()
    _pct = int(_left / MAX_FRAMES_PER_SESSION * 100)
    _cls = "pill-ok" if _pct > 40 else ("pill-warn" if _pct > 15 else "pill-err")
    st.markdown(
        f'<span class="{_cls}">🎬 {_left} analysis frames left</span>',
        unsafe_allow_html=True,
    )

    st.markdown("---")
    st.markdown("**Knowledge Base**")
    st.caption("Years of NFHS veteran officials' game notes and NFHS/MSHSL rulebook facts and interpretations")

    st.markdown("---")
    st.markdown("**Upload Files** *(Home chat)*")
    st.caption("PDFs, images, or TXT")
    chat_uploads = st.file_uploader(
        "chatfiles", type=["pdf", "jpg", "jpeg", "png", "txt"],
        accept_multiple_files=True, label_visibility="collapsed",
    )
    if chat_uploads:
        proc, names = [], []
        for uf in chat_uploads:
            c = prepare_file_content(uf)
            if c:
                proc.append(c); names.append(uf.name)
            else:
                st.warning(f"Unsupported: {uf.name}")
        st.session_state.uploaded_files_content = proc
        if names:
            st.markdown(f'<span class="pill-ok">✅ {len(names)} file(s)</span>', unsafe_allow_html=True)
    else:
        st.session_state.uploaded_files_content = []

    st.markdown("---")
    st.markdown("**Ref Log**")
    if st.session_state.messages:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        _md = chat_log_markdown()

        # 2x2 grid — PDF/Word on top (most readable), TXT/JSON below
        sb_r1c1, sb_r1c2 = st.columns(2)
        with sb_r1c1:
            _pdf = markdown_to_pdf_bytes(_md, "RefBuddy Chat Log")
            if _pdf:
                st.download_button("⬇️ PDF", data=_pdf,
                                   file_name=f"refbuddy_bb_chat_{ts}.pdf",
                                   mime="application/pdf",
                                   use_container_width=True)
        with sb_r1c2:
            _docx = markdown_to_docx_bytes(_md, "RefBuddy Chat Log")
            if _docx:
                st.download_button("⬇️ Word", data=_docx,
                                   file_name=f"refbuddy_bb_chat_{ts}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   use_container_width=True)

        sb_r2c1, sb_r2c2 = st.columns(2)
        with sb_r2c1:
            st.download_button("⬇️ TXT", data=_md,
                               file_name=f"refbuddy_bb_chat_{ts}.txt",
                               mime="text/plain", use_container_width=True)
        with sb_r2c2:
            st.download_button("⬇️ JSON", data=chat_log_json(),
                               file_name=f"refbuddy_bb_chat_{ts}.json",
                               mime="application/json", use_container_width=True)

    if st.button("🗑️ Clear Home Chat", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

    st.markdown("---")


# =============================================================================
# TABS
# =============================================================================

tab_home, tab_film, tab_ah, tab_quiz = st.tabs([
    "🏀 Home",
    "🎬 Film & Grade",
    "👥 Ref Hub",
    "📝 Quiz",
])


# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — HOME
# ─────────────────────────────────────────────────────────────────────────────

with tab_home:
    st.markdown("""
    <div class="home-hero">
        <div class="home-hero-title">🏀 RefBuddy</div>
        <div class="home-hero-slogan">Built by a ref, for refs</div>
    </div>
    """, unsafe_allow_html=True)

    chips = ["NFHS Rule Citations", "MSHSL Mods", "Shot Clock 35/20",
             "Lead · Trail · Center", "Restricted Arc", "Mercy Rule", "Photo Analysis"]
    chip_html = " &nbsp; ".join(f'<span class="pill-blue">{c}</span>' for c in chips)
    st.markdown(f'<div style="text-align:center;margin-bottom:1.4rem;line-height:2.6;">'
                f'{chip_html}</div>', unsafe_allow_html=True)

    # Quick-start prompts
    if not st.session_state.messages:
        st.markdown(f'<p style="text-align:center;color:{MUTED};font-size:0.9rem;'
                    f'margin-bottom:0.8rem;"><em>Try one of these or type your own below</em></p>',
                    unsafe_allow_html=True)
        starter_qs = [
            "What is the MSHSL bonus free throw system — how is it different from NFHS?",
            "When is offensive goaltending a violation in 2025-26?",
            "Walk me through the shot clock reset when the offense gets the rebound.",
            "What is the MSHSL mercy rule and when does running clock start?",
            "Explain the restricted area arc rule and secondary defender definition.",
            "What are Lead, Trail, and Center responsibilities on free throws?",
            "When does MSHSL's closely guarded 5-second rule apply vs. not apply?",
            "What are the 2025-26 NFHS basketball rule changes I need to know?",
        ]
        c1, c2 = st.columns(2)
        for i, q in enumerate(starter_qs):
            col = c1 if i < 4 else c2
            with col:
                if st.button(f"➤ {q}", key=f"hq_{i}", use_container_width=True):
                    st.session_state.messages.append({
                        "role": "user", "content": q,
                        "timestamp": datetime.datetime.now().isoformat(),
                    })
                    st.rerun()

    # Render chat history
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"], avatar="🏀" if msg["role"] == "user" else "⚡"):
            st.markdown(msg["content"])

    # Stream assistant reply
    if (st.session_state.messages
            and st.session_state.messages[-1]["role"] == "user"):
        if not api_key_ok():
            st.warning("⚠️ Enter your Anthropic API key in the sidebar.")
        else:
            client = make_client()
            with st.chat_message("assistant", avatar="⚡"):
                ph = st.empty()
                full = ""
                try:
                    for chunk in stream_chat(
                        client,
                        st.session_state.messages,
                        st.session_state.uploaded_files_content,
                    ):
                        full += chunk
                        ph.markdown(full + "▌")
                    ph.markdown(full)
                    st.session_state.messages.append({
                        "role": "assistant", "content": full,
                        "timestamp": datetime.datetime.now().isoformat(),
                    })
                    # ── DO NOT REMOVE ────────────────────────────────────────
                    # st.download_button bakes its `data` argument in at
                    # CREATION time. The sidebar is built earlier in the script
                    # than this Home tab, so on the run that produces an answer
                    # the sidebar's export buttons were created BEFORE this line
                    # ran — capturing a transcript containing only the question.
                    # Without this rerun, every exported log is missing the
                    # answer. The rerun rebuilds the sidebar against the now
                    # complete transcript.
                    st.rerun()
                except Exception as e:
                    st.error(handle_api_error(e))

    # Chat input pinned to bottom
    user_in = st.chat_input(
        "Ask RefBuddy a question or upload files using the "
        "Upload section in the left sidebar",
    )
    if user_in:
        st.session_state.messages.append({
            "role": "user", "content": user_in,
            "timestamp": datetime.datetime.now().isoformat(),
        })
        st.rerun()

    # Ref Log expander
    if st.session_state.messages:
        st.markdown("---")
        with st.expander("📋 Ref Log — Session Summary", expanded=False):
            _first = _human_ts(st.session_state.messages[0].get("timestamp", ""))
            _last = _human_ts(st.session_state.messages[-1].get("timestamp", ""))
            st.markdown(f"""<div class="ref-log">
            <strong>Session Stats</strong><br>
            Messages: {len(st.session_state.messages)}<br>
            Started: {_first}<br>
            Last: {_last}
            </div>""", unsafe_allow_html=True)
            for i, m in enumerate(st.session_state.messages):
                icon = "🏀 You" if m["role"] == "user" else "⚡ RefBuddy"
                st.markdown(f"**{icon}** _{_human_ts(m.get('timestamp',''))}_")
                st.markdown(m["content"])
                if i < len(st.session_state.messages) - 1:
                    st.markdown("---")

            st.markdown("---")
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            _md = chat_log_markdown()
            rl_pdf, rl_docx, rl_txt, rl_json = st.columns(4)
            with rl_pdf:
                _p = markdown_to_pdf_bytes(_md, "RefBuddy Ref Log")
                if _p:
                    st.download_button("⬇️ PDF", data=_p,
                                       file_name=f"refbuddy_bb_reflog_{ts}.pdf",
                                       mime="application/pdf",
                                       use_container_width=True)
            with rl_docx:
                _d = markdown_to_docx_bytes(_md, "RefBuddy Ref Log")
                if _d:
                    st.download_button("⬇️ Word", data=_d,
                                       file_name=f"refbuddy_bb_reflog_{ts}.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       use_container_width=True)
            with rl_txt:
                st.download_button("⬇️ TXT", data=_md,
                                   file_name=f"refbuddy_bb_reflog_{ts}.txt",
                                   mime="text/plain", use_container_width=True)
            with rl_json:
                st.download_button("⬇️ JSON", data=chat_log_json(),
                                   file_name=f"refbuddy_bb_reflog_{ts}.json",
                                   mime="application/json",
                                   use_container_width=True)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — FILM & GRADE  (merged: photos + video, ask-or-grade)
#
# The old Game Film and RefGrade tabs asked for the same upload and differed
# only in what they did with it, so they are one tab now. Accepts still PHOTOS
# as well as video clips. Photos are the common case — a ref can screenshot a
# play in two seconds, where a 40MB clip over LTE is a different story.
#
# Two analysis modes share the same uploaded media:
#   • Ask a Question  — free-form rules/mechanics analysis (SYSTEM_PROMPT)
#   • RefGrade        — structured scored evaluation (REFGRADE_PROMPT)
# ─────────────────────────────────────────────────────────────────────────────

with tab_film:
    st.markdown("## 🎬 Film & Grade")
    st.markdown(
        "Upload **photos of a call** or a **short video clip**, then either ask a "
        "question about it or run a full scored RefGrade evaluation. "
        "Every analysis begins with a Visibility Check so you know what was "
        "actually observed versus inferred."
    )

    # ── Step 1 — Upload ──────────────────────────────────────────────────────
    st.markdown("### Step 1 — Upload")
    st.info(
        "📸 **Photos work best** — screenshots from film, a still of the block/charge, "
        "or a photo of the play. Upload as many as you like.\n\n"
        "🎥 **Video** (.mp4/.mov) also works. Keep clips to 10–60 seconds and trim "
        "to the play. Large files can be slow or fail to upload on phone data — "
        "if a video won't go through, screenshot the key moments instead."
    )

    fg_uploads = st.file_uploader(
        "fg_upload",
        type=["jpg", "jpeg", "png", "mp4", "mov"],
        accept_multiple_files=True,
        label_visibility="collapsed",
        key="fg_uploader",
    )

    if fg_uploads:
        images = [f for f in fg_uploads
                  if f.name.lower().endswith((".jpg", ".jpeg", ".png"))]
        videos = [f for f in fg_uploads
                  if f.name.lower().endswith((".mp4", ".mov"))]

        # fps control only matters when a video is present
        fg_fps = 1.0
        if videos:
            fg_fps = st.select_slider(
                "Video sampling rate (frames per second)",
                options=[0.5, 1.0, 2.0], value=1.0,
                help="0.5 = overview | 1.0 = standard | 2.0 = fast action",
                key="fg_fps_slider",
            )
            st.caption(f"A 30s clip at {fg_fps} fps ≈ {int(30*fg_fps)} frames")

        if st.button("📥 Load Media for Analysis", use_container_width=True,
                     key="fg_load"):
            frames, labels = [], []
            with st.spinner("Processing upload…"):
                # Stills first — they keep their original order
                for img in images:
                    b64_img = image_to_frame_b64(img)
                    if b64_img:
                        frames.append(b64_img)
                        labels.append(img.name)
                    else:
                        st.warning(f"Could not read image: {img.name}")

                # Then any video frames
                for vid in videos:
                    if not OPENCV_AVAILABLE:
                        st.error("Video support requires opencv-python-headless.")
                        break
                    try:
                        suffix = ".mp4" if vid.name.lower().endswith(".mp4") else ".mov"
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                            tmp.write(vid.read()); tmp_path = tmp.name
                        vframes = extract_frames(tmp_path, fps=fg_fps)
                        os.unlink(tmp_path)
                        if not vframes:
                            st.error(f"No frames extracted from {vid.name} — "
                                     "check the file format or codec.")
                        frames.extend(vframes)
                        labels.extend([f"{vid.name} frame {i+1}"
                                       for i in range(len(vframes))])
                    except Exception as e:
                        st.error(f"❌ Could not process {vid.name}: {e}")

            if frames:
                st.session_state.fg_frames = frames
                st.session_state.fg_labels = labels
                st.session_state.fg_source = ", ".join(
                    [f.name for f in fg_uploads][:3]
                ) + ("…" if len(fg_uploads) > 3 else "")
                st.session_state.fg_is_video = bool(videos)
                st.session_state.fg_fps = fg_fps
                st.session_state.fg_result = ""
                st.success(
                    f"✅ Loaded {len(frames)} image(s) — "
                    f"{len(images)} photo(s), {len(frames)-len(images)} video frame(s)"
                )
            else:
                st.error("Nothing could be loaded from that upload.")

    # ── Steps 2-4 — only once media is loaded ────────────────────────────────
    if st.session_state.get("fg_frames"):
        frames = st.session_state.fg_frames
        n = len(frames)
        fps_u = st.session_state.get("fg_fps", 1.0)
        src_name = st.session_state.get("fg_source", "upload")

        st.markdown("---")
        st.markdown(f"**{n} image(s) loaded** from `{src_name}`")

        # Range selector — only useful with many frames
        if n > 1:
            st.markdown("### Step 2 — Select Range")
            sf, ef = st.slider("fg_range", 1, n, (1, min(n, 30)), key="fg_range_sel",
                               label_visibility="collapsed")
        else:
            sf = ef = 1
        sel = ef - sf + 1
        st.caption(f"Analyzing image {sf}–{ef} ({sel} total)")

        with st.expander(f"🔍 Preview {sel} selected", expanded=(sel <= 12)):
            cols = st.columns(4)
            for i, fb in enumerate(frames[sf-1:ef][:24]):
                with cols[i % 4]:
                    st.image(base64.b64decode(fb), caption=f"#{sf+i}",
                             use_container_width=True)

        # ── Step 3 — Choose what to do ───────────────────────────────────────
        st.markdown("### Step 3 — What do you want?")
        fg_mode = st.radio(
            "fg_mode",
            options=["❓ Ask a Question", "📊 RefGrade Evaluation"],
            horizontal=True,
            label_visibility="collapsed",
            key="fg_mode_sel",
        )

        # ══════════════════════════════════════════════════════════════════
        # MODE A — Ask a Question
        # ══════════════════════════════════════════════════════════════════
        if fg_mode == "❓ Ask a Question":
            st.markdown("**Quick presets**")
            p1, p2, p3 = st.columns(3)
            PRESETS = {
                "foul": ("Analyze this for any rule violations under NFHS rules and MSHSL "
                         "modifications. For each potential foul or violation: cite the rule "
                         "number, describe what you see using 'Frame N' format, state the "
                         "correct penalty and where the ball is put back in play, and note "
                         "which official (Lead/Trail/Center) had primary responsibility. "
                         "Begin with a VISIBILITY CHECK."),
                "mech": ("Evaluate the officiating mechanics visible here. Begin with a "
                         "VISIBILITY CHECK. For each visible official describe their "
                         "positioning against MSHSL 3-person mechanics (Lead under the "
                         "basket, Trail at the 28-foot mark, Center opposite the table), "
                         "whether the coverage area was correct, and any improvements. "
                         "Reference specific frame numbers."),
                "bc": ("Analyze this as a block/charge play. Check whether the defender had "
                       "established legal guarding position (Rule 4-23), whether a secondary "
                       "defender was in the restricted area arc (MSHSL Addendum Rules 4-38, "
                       "4-41), and whether any of the three exceptions apply — leading with "
                       "a foot/knee, offensive player stopping continuous movement, or "
                       "verticality. Begin with a VISIBILITY CHECK."),
            }
            # A preset writes straight into session_state and reruns.
            # Passing value= to a widget that already has a key is a no-op on
            # rerun — Streamlit keeps the stored widget state.
            with p1:
                if st.button("🔴 Was there a foul?", key="fg_p_foul",
                             use_container_width=True):
                    st.session_state["fg_q"] = PRESETS["foul"]
                    st.rerun()
            with p2:
                if st.button("⚙️ Check mechanics", key="fg_p_mech",
                             use_container_width=True):
                    st.session_state["fg_q"] = PRESETS["mech"]
                    st.rerun()
            with p3:
                if st.button("⚖️ Block or charge?", key="fg_p_bc",
                             use_container_width=True):
                    st.session_state["fg_q"] = PRESETS["bc"]
                    st.rerun()

            fg_q = st.text_area(
                "fg_question", height=110,
                placeholder="e.g. 'Is the secondary defender outside the restricted arc?' or "
                            "'Was my rotation correct as Lead on this drive?'",
                label_visibility="collapsed", key="fg_q",
            )

            if st.button(f"🔍  Analyze {sel} Image{'s' if sel != 1 else ''}",
                         disabled=not (fg_q or "").strip(),
                         type="primary",
                         use_container_width=True, key="fg_run_q"):
                if not spend_frames(sel):
                    st.stop()
                blocks = build_vision_content(
                    frames, sf-1, ef-1, fg_q.strip(), src_name, fps_u,
                    preamble_extra=(
                        "These may be still photos rather than sequential video frames. "
                        "Begin with a VISIBILITY CHECK listing which crew members are "
                        "clearly visible, partially visible, or not visible."
                    ),
                )
                st.markdown("---")
                st.markdown("#### ⚡ Analysis")
                ph = st.empty(); full = ""
                try:
                    with st.spinner(f"Analyzing {sel} image(s)…"):
                        for chunk in stream_vision(make_client(), blocks, SYSTEM_PROMPT):
                            full += chunk; ph.markdown(full + "▌")
                    ph.markdown(full)
                    st.session_state.fg_result = full
                except Exception as e:
                    st.error(handle_api_error(e))

        # ══════════════════════════════════════════════════════════════════
        # MODE B — RefGrade Evaluation
        # ══════════════════════════════════════════════════════════════════
        else:
            st.markdown("**Evaluation setup**")
            g1, g2 = st.columns(2)
            with g1:
                fg_scope = st.selectbox(
                    "Evaluate", options=[
                        "Full Crew (Overall)",
                        "Lead (L) — Under Basket",
                        "Trail (T) — Half Court",
                        "Center (C) — Middle"],
                    key="fg_scope")
            with g2:
                fg_crew = st.selectbox(
                    "Crew size",
                    options=["2-Person Crew", "3-Person Crew"],
                    key="fg_crew")

            fg_cats = st.multiselect(
                "Score these categories", options=[
                    "Positioning", "Call Accuracy", "Mechanics Execution",
                    "Dead-ball Officiating", "Communication / Signals",
                    "Shot Clock Administration", "Free Throw Administration"],
                default=["Positioning", "Call Accuracy", "Mechanics Execution",
                         "Communication / Signals"],
                key="fg_cats")

            fg_notes = st.text_area(
                "Context or focus (optional)", height=80,
                placeholder="e.g. 'A whistle was blown on the drive — was it correct?' "
                            "or 'Check my rotation timing as Lead.'",
                key="fg_notes")

            if st.button(f"📊  Run RefGrade — {fg_scope}",
                         disabled=not fg_cats,
                         type="primary",
                         use_container_width=True, key="fg_run_grade"):
                if not spend_frames(sel):
                    st.stop()
                cats = ", ".join(fg_cats)
                extra = f"\nContext: {fg_notes.strip()}" if fg_notes.strip() else ""
                gq = (f"Perform a RefGrade evaluation.\n"
                      f"Source: {src_name}\nScope: {fg_scope}\nCrew: {fg_crew}\n"
                      f"Score these categories: {cats}{extra}\n\n"
                      f"Use the exact RefGrade report structure from your system prompt. "
                      f"Begin with a VISIBILITY CHECK. Cite the specific mechanic or rule "
                      f"behind every score. If these are still photos rather than video, "
                      f"say so and scope your confidence accordingly.")
                blocks = build_vision_content(
                    frames, sf-1, ef-1, gq, src_name, fps_u,
                    preamble_extra=("Structured RefGrade evaluation. Visibility Check is "
                                    "the mandatory first section. These may be stills "
                                    "rather than sequential frames."),
                )
                st.markdown("---")
                st.markdown("#### 📊 RefGrade Report")
                ph = st.empty(); full = ""
                try:
                    with st.spinner("Running RefGrade… (20–90 seconds)"):
                        for chunk in stream_vision(make_client(), blocks, REFGRADE_PROMPT):
                            full += chunk; ph.markdown(full + "▌")
                    ph.markdown(full)
                    st.session_state.fg_result = full
                    st.session_state.rg_saved_logs.append({
                        "timestamp": datetime.datetime.now().isoformat(),
                        "source": src_name, "scope": fg_scope, "crew": fg_crew,
                        "range": f"{sf}-{ef}", "result": full,
                    })
                except Exception as e:
                    st.error(handle_api_error(e))

        # ── Step 4 — Export ──────────────────────────────────────────────────
        if st.session_state.get("fg_result"):
            st.markdown("---")
            e1, e2, e3 = st.columns(3)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            with e1:
                st.download_button("⬇️ Download TXT",
                                   data=st.session_state.fg_result,
                                   file_name=f"refbuddy_analysis_{ts}.txt",
                                   mime="text/plain", use_container_width=True)
            with e2:
                _pdf = markdown_to_pdf_bytes(st.session_state.fg_result,
                                             "RefBuddy Film Analysis")
                if _pdf:
                    st.download_button("⬇️ Export PDF", data=_pdf,
                                       file_name=f"refbuddy_analysis_{ts}.pdf",
                                       mime="application/pdf",
                                       use_container_width=True)
            with e3:
                if st.button("🗑️ Clear", use_container_width=True, key="fg_clear"):
                    for k in ("fg_frames", "fg_labels", "fg_source",
                              "fg_is_video", "fg_result"):
                        st.session_state.pop(k, None)
                    st.rerun()

    elif not fg_uploads:
        st.markdown("---")
        st.markdown("""<div class="rb-card-blue">
        <h4 style="margin-top:0;color:#003087;">How to Use Film &amp; Grade</h4>
        <ol style="color:#1F2937;line-height:2.0;">
        <li><b>Upload photos</b> of the play — screenshots, stills, or a phone photo.
            Video clips work too, but photos are faster and more reliable.</li>
        <li><b>Load Media</b> — everything becomes analyzable images</li>
        <li><b>Pick a mode</b> — ask a specific question, or run a scored RefGrade</li>
        <li><b>Review</b> — every answer cites NFHS rules and MSHSL modifications,
            and starts with a Visibility Check</li>
        <li><b>Export</b> to TXT or PDF</li>
        </ol>
        <p style="font-size:0.82rem;color:#4B5563;margin-bottom:0;">
        <em>Tip: three or four well-chosen stills (defender establishing position, point
        of contact, and the finish) usually beat a full video clip for both speed and
        cost.</em>
        </p>
        </div>""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — REF HUB  (Pre-Game first, Crew + Ref Eval merged)
#
# Two sections instead of three. Pre-Game Meeting leads because it is the most
# frequently used and needs no upload. Crew Eval and Ref Eval were near-identical
# forms differing only in scope, so they are now one Evaluations section with a
# scope dropdown — "Full Crew" routes to CREW_EVAL_PROMPT, any single position
# routes to REF_EVAL_PROMPT. Both accept photos as well as video.
# ─────────────────────────────────────────────────────────────────────────────

with tab_ah:
    st.markdown("## 👥 Ref Hub")
    st.markdown(
        "Auto-generated pre-game meeting agendas, plus film- and photo-based "
        "crew and individual official evaluations."
    )

    # ── Section selector ─────────────────────────────────────────────────────
    sub_a, sub_b = st.columns(2)
    with sub_a:
        _active = st.session_state.ah_sub == "pregame"
        if st.button("📅 Pre-Game Meeting", use_container_width=True,
                     key="ah_sub_pregame",
                     type="primary" if _active else "secondary"):
            st.session_state.ah_sub = "pregame"
            st.rerun()
    with sub_b:
        _active = st.session_state.ah_sub == "eval"
        if st.button("🎬 Evaluations", use_container_width=True,
                     key="ah_sub_eval",
                     type="primary" if _active else "secondary"):
            st.session_state.ah_sub = "eval"
            st.rerun()

    st.markdown("---")

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1 — PRE-GAME MEETING
    # ═══════════════════════════════════════════════════════════════════════
    if st.session_state.ah_sub == "pregame":
        st.markdown("### 📅 Pre-Game Meeting Agenda Generator")
        st.markdown(
            "Auto-generates a pre-game agenda from the RefBuddy Knowledge Base, "
            "including current rule changes, MSHSL modifications, shot clock "
            "reminders, and your own notes section for a fully customized meeting."
        )

        pg1, pg2 = st.columns(2)
        with pg1:
            pg_crew = st.selectbox("Crew size",
                                    ["2-Person Crew", "3-Person Crew"],
                                    key="pg_crew_sel")
            pg_level = st.selectbox("Game level",
                                     ["Varsity", "Junior Varsity",
                                      "9th Grade (Sub-Varsity)", "Playoff"],
                                     key="pg_level_sel")
        with pg2:
            pg_date = st.text_input("Game date (optional)",
                                     placeholder="e.g. Friday, January 16",
                                     key="pg_date")
            pg_teams = st.text_input("Teams (optional)",
                                      placeholder="e.g. Eden Prairie vs Wayzata",
                                      key="pg_teams")

        pg_focus = st.multiselect(
            "Additional emphasis topics (optional)",
            options=["2025-26 Rule Changes (No Offensive Goaltending, Backboard Slap)",
                     "Shot Clock Resets (35 vs 20 Scenarios)",
                     "Restricted Area Arc & Secondary Defender",
                     "Flopping Mechanics & Warning Protocol",
                     "Mercy Rule Procedure", "Overtime Procedure",
                     "Closely Guarded — MSHSL Dribbling Exception",
                     "Bonus Free Throw System (MSHSL 1&1 vs NFHS Quarter)",
                     "Throw-In Spots — 3-Point Line Demarcation",
                     "Uniform Compliance (Home=Dark, Visitor=White)",
                     "Free Throw Administration (C/L/T responsibilities)",
                     "Block/Charge — Restricted Area Signaling"],
            key="pg_focus_sel",
        )

        pg_assignor_notes = st.text_area(
            "Assignor's Custom Notes / Emphasis",
            height=130,
            placeholder="Add anything you want to emphasize for this game",
            key="pg_assignor_notes",
        )

        if st.button("📅  Generate Pre-Game Meeting Agenda",
                     type="primary", use_container_width=True, key="pg_generate"):
            focus_str = (f"Additional emphasis topics requested: {', '.join(pg_focus)}\n"
                         if pg_focus else "")
            header_str = ""
            if pg_date or pg_teams:
                header_str = (f"Game: {pg_teams or 'TBD'} | "
                              f"Date: {pg_date or 'TBD'} | {pg_level}\n")

            # Preserve line breaks from the notes box as separate bullets so
            # multi-line notes are not run together in the agenda
            if pg_assignor_notes.strip():
                raw_lines = [l.strip() for l in pg_assignor_notes.strip().splitlines()
                             if l.strip()]
                if len(raw_lines) == 1:
                    notes_section = raw_lines[0]
                else:
                    notes_section = "\n".join(f"- {ln.lstrip('-* ').strip()}"
                                               for ln in raw_lines)
            else:
                notes_section = "(No specific assignor notes provided for this game.)"

            prompt = (
                f"Generate a pre-game meeting agenda for the following game.\n\n"
                f"{header_str}"
                f"Crew configuration: {pg_crew}\n"
                f"Game level: {pg_level}\n"
                f"{focus_str}\n"
                f"For the Assignor Notes section, use EXACTLY this content verbatim — "
                f"do not summarize or rephrase it:\n{notes_section}\n\n"
                f"Generate the full agenda following your system prompt structure. "
                f"Make all sections specific to basketball and MSHSL."
            )

            with st.spinner("Generating pre-game meeting agenda… (15–30 seconds)"):
                try:
                    result = call_api_sync(prompt, PREGAME_MEETING_PROMPT, max_tokens=3000)
                    st.session_state.ah_pregame_result = result
                    st.session_state.ah_pregame_logs.append({
                        "timestamp": datetime.datetime.now().isoformat(),
                        "teams": pg_teams or "Unknown",
                        "date": pg_date or "Unknown",
                        "crew": pg_crew, "level": pg_level,
                        "result": result,
                    })
                    st.success("✅ Agenda generated!")
                except Exception as e:
                    st.error(handle_api_error(e))

        if st.session_state.ah_pregame_result:
            st.markdown("---")
            with st.expander("📋 Pre-Game Meeting Agenda", expanded=True):
                st.markdown(st.session_state.ah_pregame_result)

            st.markdown("**Export**")
            ep1, ep2, ep3, ep4 = st.columns(4)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            with ep1:
                st.download_button("⬇️ TXT",
                                   data=st.session_state.ah_pregame_result,
                                   file_name=f"pregame_{ts}.txt",
                                   mime="text/plain", use_container_width=True)
            with ep2:
                pdf_b = markdown_to_pdf_bytes(st.session_state.ah_pregame_result,
                                               "Pre-Game Meeting Agenda — RefBuddy")
                if pdf_b:
                    st.download_button("⬇️ PDF",
                                       data=pdf_b, file_name=f"pregame_{ts}.pdf",
                                       mime="application/pdf", use_container_width=True)
                else:
                    st.caption("pip install fpdf2")
            with ep3:
                docx_b = markdown_to_docx_bytes(st.session_state.ah_pregame_result,
                                                 "Pre-Game Meeting Agenda — RefBuddy")
                if docx_b:
                    st.download_button("⬇️ Word",
                                       data=docx_b, file_name=f"pregame_{ts}.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       use_container_width=True)
                else:
                    st.caption("pip install python-docx")
            with ep4:
                if st.button("🗑️ Clear", use_container_width=True, key="pg_clear"):
                    st.session_state.ah_pregame_result = ""; st.rerun()

            if st.session_state.ah_pregame_logs:
                st.markdown("---")
                st.markdown(f"**Agenda History ({len(st.session_state.ah_pregame_logs)} saved)**")
                for log in reversed(st.session_state.ah_pregame_logs[-5:]):
                    st.markdown(
                        f'<div class="rb-card" style="padding:0.7rem 1rem;">'
                        f'<strong>{log["teams"]}</strong> — {log["date"]} | {log["level"]} | {log["crew"]}'
                        f'<br><span style="font-size:0.78rem;color:{MUTED};">{log["timestamp"][:19]}</span>'
                        f'</div>', unsafe_allow_html=True
                    )
                ts2 = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    f"⬇️ Download All Agendas ({len(st.session_state.ah_pregame_logs)})",
                    data=json.dumps(st.session_state.ah_pregame_logs, indent=2, ensure_ascii=False),
                    file_name=f"pregame_all_{ts2}.json",
                    mime="application/json", use_container_width=True,
                )

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2 — EVALUATIONS  (crew + individual merged behind a scope picker)
    # ═══════════════════════════════════════════════════════════════════════
    else:
        st.markdown("### 🎬 Evaluations")
        st.markdown(
            "Upload photos or film and evaluate either the whole crew or one "
            "specific official. Every report opens with a Visibility Check and "
            "cites NFHS rules and MSHSL mechanics throughout."
        )

        if not OPENCV_AVAILABLE:
            st.warning("Video support requires opencv-python-headless. "
                       "Photo uploads still work.")

        ev_uploads = st.file_uploader(
            "ev_upload",
            type=["jpg", "jpeg", "png", "mp4", "mov"],
            accept_multiple_files=True,
            label_visibility="collapsed",
            key="ah_ev_uploader",
        )
        st.caption("📸 Photos (JPG/PNG) or 🎥 video (MP4/MOV). Photos upload faster "
                   "and are usually enough — try the defender establishing position, "
                   "the point of contact, and the finish.")

        e1, e2 = st.columns(2)
        with e1:
            ev_scope = st.selectbox(
                "Evaluate",
                options=["Full Crew (Overall)",
                         "Lead (L) — Under Basket",
                         "Trail (T) — Half Court",
                         "Center (C) — Middle"],
                key="ev_scope",
            )
        with e2:
            ev_crew = st.selectbox(
                "Crew configuration",
                ["2-Person Crew", "3-Person Crew"],
                key="ev_crew",
            )

        ev_fps = 1.0
        if ev_uploads and any(f.name.lower().endswith((".mp4", ".mov")) for f in ev_uploads):
            ev_fps = st.select_slider(
                "Video sampling rate (frames per second)",
                options=[0.5, 1.0, 2.0], value=1.0, key="ev_fps",
            )

        ev_notes = st.text_area(
            "Notes / focus (optional)", height=90,
            placeholder="e.g. 'Check Lead rotation timing on the drive.' or "
                        "'A whistle was blown here — was it correct?'",
            key="ev_notes",
        )

        is_crew = ev_scope.startswith("Full Crew")
        btn_label = ("🎬  Generate Crew Evaluation" if is_crew
                     else f"🏀  Generate {ev_scope} Evaluation")

        if st.button(btn_label, type="primary", use_container_width=True,
                     disabled=not ev_uploads, key="ev_run"):
            frames = []
            with st.spinner("Processing upload…"):
                for f in ev_uploads:
                    if f.name.lower().endswith((".jpg", ".jpeg", ".png")):
                        b = image_to_frame_b64(f)
                        if b:
                            frames.append(b)
                    else:
                        if not OPENCV_AVAILABLE:
                            st.error("Video requires opencv-python-headless.")
                            break
                        try:
                            sfx = ".mp4" if f.name.lower().endswith(".mp4") else ".mov"
                            with tempfile.NamedTemporaryFile(delete=False, suffix=sfx) as t:
                                t.write(f.read()); tp = t.name
                            frames.extend(extract_frames(tp, fps=ev_fps))
                            os.unlink(tp)
                        except Exception as e:
                            st.error(f"❌ Could not process {f.name}: {e}")

            if not frames:
                st.error("Nothing could be read from that upload.")
            else:
                cap = min(len(frames), 40)
                if spend_frames(cap):
                    name = ", ".join(f.name for f in ev_uploads[:3])
                    extra = f"\nNotes: {ev_notes.strip()}" if ev_notes.strip() else ""
                    if is_crew:
                        q = (f"Perform a full crew evaluation.\nSource: {name}\n"
                             f"Crew configuration: {ev_crew}\n"
                             f"Images analyzed: 1–{cap} of {len(frames)}{extra}\n\n"
                             f"Analyze all visible officials (Lead/Trail/Center) for "
                             f"positioning, call accuracy, mechanics, dead-ball "
                             f"officiating, shot clock administration and "
                             f"communication. Begin with a VISIBILITY CHECK.")
                        prompt_used = CREW_EVAL_PROMPT
                        heading = "📊 Crew Evaluation Report"
                    else:
                        q = (f"Evaluate ONLY the {ev_scope} in this material.\n"
                             f"Source: {name}\nCrew configuration: {ev_crew}\n"
                             f"Images analyzed: 1–{cap} of {len(frames)}{extra}\n\n"
                             f"Focus entirely on this one official. Ignore others unless "
                             f"their actions directly affect this official's "
                             f"responsibilities. Begin with a VISIBILITY CHECK for this "
                             f"position only.")
                        prompt_used = REF_EVAL_PROMPT
                        heading = f"📊 {ev_scope} Evaluation Report"

                    blocks = build_vision_content(
                        frames, 0, cap - 1, q, name, ev_fps,
                        preamble_extra=("These may be still photos rather than "
                                        "sequential video frames. Visibility Check "
                                        "is the mandatory first section."),
                    )
                    st.markdown("---")
                    st.markdown(f"#### {heading}")
                    ph = st.empty(); full = ""
                    try:
                        with st.spinner(f"Analyzing {cap} image(s)… (30–120 seconds)"):
                            for chunk in stream_vision(make_client(), blocks, prompt_used):
                                full += chunk; ph.markdown(full + "▌")
                        ph.markdown(full)
                        st.session_state.ah_eval_result = full
                        st.session_state.ah_eval_scope = ev_scope
                    except Exception as e:
                        st.error(handle_api_error(e))

        if st.session_state.get("ah_eval_result"):
            st.markdown("---")
            with st.expander("📄 Evaluation Report", expanded=True):
                st.markdown(st.session_state.ah_eval_result)

            st.markdown("**Export**")
            z1, z2, z3 = st.columns(3)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            with z1:
                st.download_button("⬇️ TXT", data=st.session_state.ah_eval_result,
                                   file_name=f"evaluation_{ts}.txt",
                                   mime="text/plain", use_container_width=True)
            with z2:
                _p = markdown_to_pdf_bytes(
                    st.session_state.ah_eval_result,
                    f"{st.session_state.get('ah_eval_scope','')} Evaluation Report")
                if _p:
                    st.download_button("⬇️ PDF", data=_p,
                                       file_name=f"evaluation_{ts}.pdf",
                                       mime="application/pdf",
                                       use_container_width=True)
            with z3:
                if st.button("🗑️ Clear", use_container_width=True, key="ev_clear"):
                    st.session_state.ah_eval_result = ""
                    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# TAB 5 — QUIZ & DRILLS
# ─────────────────────────────────────────────────────────────────────────────

with tab_quiz:
    st.markdown("## 📝 Quiz")
    st.markdown("Test your basketball officiating knowledge! Questions are generated "
                "from the RefBuddy Knowledge Base with a 50/50 mix of multiple-choice "
                "and true/false.")

    # Mode selector
    mode_c1, mode_c2 = st.columns(2)
    with mode_c1:
        active_flash = st.session_state.quiz_mode == "flashcard"
        if st.button(
            "⚡ Flashcard Mode\n\nOne question at a time — instant feedback",
            use_container_width=True, key="q_mode_flash",
            type="primary" if active_flash else "secondary",
        ):
            st.session_state.quiz_mode = "flashcard"
            st.session_state.quiz_current_q = None
            st.session_state.quiz_answered = False
            st.rerun()

    with mode_c2:
        active_10 = st.session_state.quiz_mode == "ten_questions"
        if st.button(
            "📋 10-Question Quiz\n\nFull quiz with score report and review",
            use_container_width=True, key="q_mode_10",
            type="primary" if active_10 else "secondary",
        ):
            st.session_state.quiz_mode = "ten_questions"
            st.session_state.tenq_questions = []
            st.session_state.tenq_index = 0
            st.session_state.tenq_answers = []
            st.session_state.tenq_finished = False
            st.session_state.tenq_answered_this = False
            st.session_state.tenq_user_answer = None
            st.rerun()

    st.markdown("---")

    # Topic selector
    quiz_topics = ["Mixed", "Rules (NFHS)", "MSHSL Specific", "Shot Clock",
                   "Mechanics & Positioning", "2025-26 Changes", "Restricted Area Arc",
                   "Fouls & Free Throws", "Game Situations"]
    if st.session_state.quiz_mode:
        st.session_state.quiz_topic = st.selectbox(
            "Topic Focus", quiz_topics, key="quiz_topic_sel"
        )

    # ── FLASHCARD MODE ─────────────────────────────────────────────────────────

    if st.session_state.quiz_mode == "flashcard":
        if st.session_state.quiz_total > 0:
            accuracy_display(st.session_state.quiz_correct, st.session_state.quiz_total)

        btn_label = "🔄 Next Question" if st.session_state.quiz_current_q else "🎯 Get Question"
        if st.button(btn_label, use_container_width=True, key="get_q_btn"):
            with st.spinner("Generating question…"):
                q = generate_single_question(
                    st.session_state.quiz_topic,
                    st.session_state.quiz_session_topics,
                )
            if q:
                st.session_state.quiz_current_q = q
                st.session_state.quiz_answered = False
                st.session_state.quiz_user_answer = None
                topic_tag = q.get("topic", "")
                if topic_tag:
                    st.session_state.quiz_session_topics.append(topic_tag)
                st.rerun()

        if st.session_state.quiz_current_q:
            q = st.session_state.quiz_current_q
            render_question_card(q)

            if not st.session_state.quiz_answered:
                options = q.get("options", {})
                option_labels = [f"{k}:  {v}" for k, v in sorted(options.items())]
                user_choice = st.radio("**Select your answer:**", option_labels, key="quiz_radio")
                if st.button("✅ Submit Answer", use_container_width=True, key="submit_q"):
                    chosen = user_choice.split(":")[0].strip()
                    st.session_state.quiz_user_answer = chosen
                    st.session_state.quiz_answered = True
                    is_correct = render_feedback(q, chosen)
                    st.session_state.quiz_total += 1
                    if is_correct:
                        st.session_state.quiz_correct += 1
                    st.rerun()
            else:
                render_feedback(q, st.session_state.quiz_user_answer)

        if st.session_state.quiz_total > 0:
            st.markdown("---")
            rc1, rc2 = st.columns(2)
            with rc1:
                if st.button("🗑️ Reset Score", use_container_width=True, key="reset_score"):
                    st.session_state.quiz_total = 0
                    st.session_state.quiz_correct = 0
                    st.session_state.quiz_current_q = None
                    st.session_state.quiz_answered = False
                    st.session_state.quiz_session_topics = []
                    st.rerun()

    # ── 10-QUESTION MODE ───────────────────────────────────────────────────────

    elif st.session_state.quiz_mode == "ten_questions":
        if not st.session_state.tenq_questions and not st.session_state.tenq_finished:
            if st.button("🎯 Generate 10-Question Quiz", use_container_width=True, key="gen_10q"):
                with st.spinner("Generating 10 basketball rules questions… (20–40 seconds)"):
                    qs = generate_ten_questions(st.session_state.quiz_topic)
                if qs:
                    st.session_state.tenq_questions = qs
                    st.session_state.tenq_index = 0
                    st.session_state.tenq_answers = []
                    st.session_state.tenq_finished = False
                    st.session_state.tenq_answered_this = False
                    st.session_state.tenq_user_answer = None
                    st.rerun()

        elif st.session_state.tenq_questions and not st.session_state.tenq_finished:
            questions = st.session_state.tenq_questions
            total_qs = len(questions)
            idx = st.session_state.tenq_index

            st.markdown(f'<div style="text-align:right;color:{MUTED};font-size:0.88rem;margin-bottom:0.5rem;">'
                        f'Question {idx+1} of {total_qs} &nbsp;|&nbsp; Topic: {st.session_state.quiz_topic}</div>',
                        unsafe_allow_html=True)

            progress_pct = idx / total_qs
            st.markdown(f"""
            <div class="accuracy-bar-wrap">
                <div class="accuracy-bar-fill" style="width:{int(progress_pct*100)}%;"></div>
            </div>""", unsafe_allow_html=True)

            if idx < total_qs:
                q = questions[idx]
                options = q.get("options", {})
                render_question_card(q, question_num=f"Q{idx+1}.")

                if not st.session_state.tenq_answered_this:
                    option_labels = [f"{k}:  {v}" for k, v in sorted(options.items())]
                    user_choice = st.radio("**Select your answer:**", option_labels,
                                           key=f"tenq_radio_{idx}")
                    if st.button("✅ Submit Answer", use_container_width=True,
                                 key=f"tenq_submit_{idx}"):
                        chosen = user_choice.split(":")[0].strip()
                        st.session_state.tenq_user_answer = chosen
                        st.session_state.tenq_answered_this = True
                        is_correct = chosen == q.get("correct", "")
                        st.session_state.tenq_answers.append({
                            "question_num": idx + 1,
                            "user": chosen, "correct": q.get("correct", ""),
                            "is_correct": is_correct, "data": q,
                        })
                        st.rerun()
                else:
                    render_feedback(q, st.session_state.tenq_user_answer)
                    st.markdown("")
                    is_last = (idx == total_qs - 1)
                    btn_lbl = "📊 See Final Score" if is_last else f"➡️ Next ({idx+2}/{total_qs})"
                    if st.button(btn_lbl, use_container_width=True, key=f"tenq_next_{idx}"):
                        if is_last:
                            st.session_state.tenq_finished = True
                        else:
                            st.session_state.tenq_index += 1
                            st.session_state.tenq_answered_this = False
                            st.session_state.tenq_user_answer = None
                        st.rerun()

        elif st.session_state.tenq_finished and st.session_state.tenq_answers:
            answers = st.session_state.tenq_answers
            n_correct = sum(1 for a in answers if a["is_correct"])
            n_total = len(answers)
            pct = int(round(n_correct / n_total * 100))
            score_color = ("#15803D" if pct >= 80 else ("#92400E" if pct >= 60 else "#991B1B"))
            grade_label = ("🏆 Excellent!" if pct >= 90 else "✅ Good" if pct >= 80
                           else "📈 Getting there" if pct >= 70 else "📚 Keep studying"
                           if pct >= 60 else "🔁 Review the material")

            st.markdown(f"""
            <div style="background:{CARD};border:2px solid {score_color};border-radius:14px;
                        padding:2rem;text-align:center;margin-bottom:1.5rem;
                        box-shadow:0 4px 16px rgba(0,0,0,0.08);">
                <div style="font-size:3.5rem;font-weight:900;color:{score_color};">{pct}%</div>
                <div style="font-size:1.3rem;font-weight:700;color:#1F2937;margin:0.3rem 0;">
                    {n_correct} / {n_total} correct &nbsp; {grade_label}</div>
                <div style="color:{MUTED};font-size:0.9rem;">Topic: {st.session_state.quiz_topic}</div>
            </div>""", unsafe_allow_html=True)

            ra1, ra2 = st.columns(2)
            with ra1:
                if st.button("📁 Save Results to My Log", use_container_width=True, key="tenq_save"):
                    st.session_state.quiz_log.append({
                        "timestamp": datetime.datetime.now().isoformat(),
                        "topic": st.session_state.quiz_topic,
                        "score": pct, "correct": n_correct, "total": n_total,
                        "answers": answers,
                    })
                    st.success(f"✅ Saved! {len(st.session_state.quiz_log)} quiz log(s) on file.")
            with ra2:
                if st.button("🔄 Take Another Quiz", use_container_width=True, key="tenq_restart"):
                    st.session_state.tenq_questions = []
                    st.session_state.tenq_index = 0
                    st.session_state.tenq_answers = []
                    st.session_state.tenq_finished = False
                    st.session_state.tenq_answered_this = False
                    st.session_state.tenq_user_answer = None
                    st.rerun()

            st.markdown("---")
            st.markdown("### 📋 Full Review")
            for a in answers:
                qd = a["data"]
                opts = qd.get("options", {})
                u, c, ic = a["user"], a["correct"], a["is_correct"]
                icon = "✅" if ic else "❌"
                cbg = "#F0FDF4" if ic else "#FFF1F2"
                cbo = "#4ADE80" if ic else "#F87171"
                u_txt, c_txt = opts.get(u, u), opts.get(c, c)
                corr_line = (
                    "" if ic
                    else f'<br><strong style="color:#7F1D1D;">✔ Correct: {c}: {c_txt}</strong>'
                )
                st.markdown(f"""
                <div style="background:{cbg};border:1.5px solid {cbo};border-radius:10px;
                            padding:1.1rem 1.3rem;margin-bottom:0.9rem;">
                    <div style="font-weight:700;color:#1F2937;">
                        {icon} Q{a["question_num"]}: {qd.get("question","")}</div>
                    <div style="font-size:0.9rem;color:#1F2937;margin-top:0.3rem;">
                        <strong>Your answer:</strong> {u}: {u_txt}{corr_line}</div>
                </div>""", unsafe_allow_html=True)
                with st.expander(f"📖 Explanation — Q{a['question_num']}", expanded=False):
                    p = qd.get("personal_note", "")
                    pnote = f'<br><strong>📋 From RefBuddy Knowledge Base:</strong> {p}' if p else ""
                    st.markdown(f"""<div class="quiz-explanation">
                    {qd.get("explanation","")}<br><br>
                    <strong>📌 Citation:</strong> {qd.get("rule_citation","")}{pnote}
                    </div>""", unsafe_allow_html=True)

            if st.session_state.quiz_log:
                st.markdown("---")
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    f"⬇️ Download All Quiz Results ({len(st.session_state.quiz_log)} saved)",
                    data=json.dumps(st.session_state.quiz_log, indent=2, ensure_ascii=False),
                    file_name=f"refbuddy_bb_quiz_{ts}.json",
                    mime="application/json", use_container_width=True,
                )

    elif st.session_state.quiz_mode is None:
        st.markdown("""<div class="rb-card" style="text-align:center;padding:1.5rem;">
        <p style="color:#4B5563;margin:0;">👆 Select a mode above to get started.</p>
        </div>""", unsafe_allow_html=True)


# =============================================================================
# FOOTER
# =============================================================================

st.markdown(f"""
<div class="rb-footer">
    Built by a ref, for refs 🏀 &nbsp;|&nbsp;
    Years of NFHS veteran officials&rsquo; game notes and NFHS/MSHSL rulebook facts and interpretations<br>
    <span style="font-size:0.72rem;">
    Always confirm rulings with your MSHSL assignor. Not official NFHS/MSHSL interpretation.
    </span>
</div>
""", unsafe_allow_html=True)

# ── Terms of Use / Disclaimer ────────────────────────────────────────────────
with st.expander("📄 Terms of Use, Disclaimer & Attribution", expanded=False):
    st.markdown("""
#### Not an official source
RefBuddy is an independent study and preparation tool built by a working MSHSL
basketball official. **It is not affiliated with, endorsed by, sponsored by, or
approved by the National Federation of State High School Associations (NFHS),
the Minnesota State High School League (MSHSL), or any officials' association.**
NFHS and MSHSL names are used here only to identify the rule sets being discussed.

#### No rulings, no substitute for the rulebook
Nothing in RefBuddy is an official rule interpretation. Answers are generated by
an AI model and **can be wrong, incomplete, or out of date.** RefBuddy summarizes
rules in plain language and points you to rule numbers — it does not reproduce
rulebook text. **You are expected to own and consult a current NFHS Basketball
Rules Book and the current MSHSL Minnesota modifications.** For any ruling that
matters, confirm with your district assignor or state rules interpreter before
acting on it.

#### Intellectual property
Rule numbers, penalty values, timing values, shot clock reset triggers, and
mechanics positions are factual information. The NFHS Rules Book, Case Book,
Points of Emphasis, and MSHSL publications are copyrighted works owned by their
respective organizations. RefBuddy holds condensed facts, rule citations, and
limited definitional excerpts for private study by password-holding officials;
it does not distribute the publications themselves, and you are expected to own
current editions. Personal game notes and
all original commentary are the author's own work.

#### Your uploads
Photos, video clips, and files you upload are held in memory for the length of
your session, sent to the Anthropic API for analysis, and are not stored on a
server by this app. Do not upload anything containing sensitive personal
information about students, players, coaches, or officials. Chat logs and reports
you download are saved only to your own device.

#### Acceptable use
Access is limited to officials who have been given the crew password. Don't
share it publicly. Usage is capped per session to keep costs sustainable.

#### No warranty
This tool is provided "as is," without warranty of any kind. The author is not
liable for any outcome resulting from its use, including missed or incorrect
calls, evaluations, or game administration decisions.
    """)
